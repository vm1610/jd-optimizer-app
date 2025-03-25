import streamlit as st
import boto3
import json
import pandas as pd
import numpy as np
import os
from docx import Document
import re
from collections import Counter
import plotly.graph_objects as go
import plotly.express as px
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import nltk
import datetime
import time
from typing import Dict, List, Any, Optional
import uuid
from jdoptim_logger import JDOptimLogger
anthropic_client='None'



# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')

# Initialize lemmatizer and stopwords
lemmatizer = WordNetLemmatizer()
stop_words = set(stopwords.words('english'))

# Custom CSS for styling
custom_css = """
<style>
    .main-header {
        font-size: 24px;
        font-weight: bold;
        margin-bottom: 15px;
        color: #1E3A8A;
    }
    
    .section-header {
        font-size: 20px;
        font-weight: bold;
        margin-top: 25px;
        margin-bottom: 10px;
        color: #1E3A8A;
    }
    
    .subsection-header {
        font-size: 16px;
        font-weight: bold;
        margin-top: 15px;
        margin-bottom: 5px;
        color: #2563EB;
    }
    
    .highlight-box {
        background-color: #F3F4F6;
        border-left: 4px solid #2563EB;
        padding: 10px;
        margin: 10px 0;
    }
    
    .success-box {
        background-color: #D1FAE5;
        border-left: 4px solid #10B981;
        padding: 10px;
        margin: 10px 0;
    }
    
    .warning-box {
        background-color: #FEF3C7;
        border-left: 4px solid #F59E0B;
        padding: 10px;
        margin: 10px 0;
    }
    
    .tab-header {
        font-weight: bold;
        color: #2563EB;
    }
    
    .metric-card {
        background-color: #f8f9fa;
        border-radius: 5px;
        padding: 10px;
        margin: 5px 0;
        box-shadow: 1px 1px 5px rgba(0,0,0,0.1);
    }
    
    .category-high {
        background-color: #e6ffe6;
        border-left: 3px solid #2ecc71;
        padding: 10px;
        margin: 5px 0;
        border-radius: 5px;
    }
    
    .category-medium {
        background-color: #fff5e6;
        border-left: 3px solid #f39c12;
        padding: 10px;
        margin: 5px 0;
        border-radius: 5px;
    }
    
    .category-low {
        background-color: #ffe6e6;
        border-left: 3px solid #e74c3c;
        padding: 10px;
        margin: 5px 0;
        border-radius: 5px;
    }
    
    div[data-testid="stSidebarNav"] {
        background-color: #F3F4F6;
        padding-top: 2rem;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 40px;
        white-space: pre-wrap;
        background-color: #F9FAFB;
        border-radius: 4px 4px 0 0;
        gap: 1px;
        padding-top: 8px;
        padding-bottom: 8px;
    }
    
    .stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
        font-size: 14px;
        font-weight: 500;
    }
    
    .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {
        background-color: #DBEAFE;
    }
</style>
"""

def read_job_description(file_path):
    """Read job description from either .txt or .docx file"""
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")

# Function to save enhanced JD
def save_enhanced_jd(content, filename, format_type):
    """Save job description content to a file"""
    if format_type == 'docx':
        doc = Document()
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(filename)
        return True
    elif format_type == 'txt':
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(content)
        return True
    return False

class JobDescriptionAnalyzer:
    """Analyzes job descriptions for skill coverage and other metrics"""
    def __init__(self):
        self.categories = {
            'Technical Skills': ['python', 'java', 'javascript', 'c#', 'ruby', 'sql', 'aws', 'azure', 'cloud', 
                               'docker', 'kubernetes', 'api', 'database', 'git', 'linux', 'agile', 'devops', 
                               'ml', 'ai', 'analytics', 'full-stack', 'frontend', 'backend'],
            'Soft Skills': ['communication', 'leadership', 'teamwork', 'collaboration', 'problem-solving', 
                           'analytical', 'initiative', 'organizational', 'time management', 'interpersonal'],
            'Experience Level': ['year', 'years', 'senior', 'junior', 'mid-level', 'lead', 'manager', 'experience'],
            'Education': ['degree', 'bachelor', 'master', 'phd', 'certification', 'education'],
            'Tools & Technologies': ['jira', 'confluence', 'slack', 'github', 'gitlab', 'azure', 'jenkins', 
                                   'terraform', 'react', 'angular', 'vue', 'node'],
            'Domain Knowledge': ['finance', 'healthcare', 'retail', 'banking', 'insurance', 'technology', 
                               'manufacturing', 'telecom', 'e-commerce']
        }

    def analyze_text(self, text):
        """Analyze text for keyword coverage in each category"""
        if not text:
            return {category: 0.0 for category in self.categories}
            
        text = text.lower()
        scores = {}
        
        for category, keywords in self.categories.items():
            category_score = 0
            for keyword in keywords:
                count = len(re.findall(r'\b' + keyword + r'\b', text))
                category_score += count
            max_possible = len(keywords)
            scores[category] = min(category_score / max_possible, 1.0)
            
        return scores

class JobDescriptionAgent:
    """Agent for enhancing job descriptions using AWS Bedrock Claude"""
    def __init__(self, model_id, max_tokens=5000, temperature=0.7):
        self.model_id = model_id
        self.max_tokens = max_tokens
        self.temperature = temperature
        
        try:
            # SECURITY: Replace hardcoded credentials with proper credential management
            self.client = boto3.client(
                service_name='bedrock-runtime',
                aws_access_key_id=st.secrets["aws"]["access_key"],
                aws_secret_access_key=st.secrets["aws"]["secret_key"],
                region_name=st.secrets["aws"]["region"],
            )
        except Exception as e:
            st.error(f"AWS Bedrock client initialization failed: {e}")
            print(f"Error initializing AWS Bedrock client: {e}")
            self.client = None
    def generate_initial_descriptions(self, job_description):
        """Generate detailed and structured job descriptions based on the given job description."""
        # If client is not initialized properly, return dummy versions
        if not self.client:
            return [
                f"Enhanced Version 1 (Example):\n\nOverview: This role is responsible for...\n\nKey Responsibilities:\n- Responsibility 1\n- Responsibility 2\n\nRequired Skills:\n- Skill 1\n- Skill 2",
                f"Enhanced Version 2 (Example):\n\nOverview: This position focuses on...\n\nKey Responsibilities:\n- Responsibility A\n- Responsibility B\n\nRequired Skills:\n- Skill A\n- Skill B",
                f"Enhanced Version 3 (Example):\n\nOverview: A key position that...\n\nKey Responsibilities:\n- Primary task 1\n- Primary task 2\n\nRequired Skills:\n- Critical skill 1\n- Critical skill 2"
            ]
        
        prompt = (
            "You are a job description specialist. Your task is to refine and expand upon the provided job description, "
            "creating three distinct versions that are structured, detailed, and aligned with industry best practices.\n\n"
            
            "### Guidelines:\n"
            "- Do NOT make assumptions or introduce inaccuracies.\n"
            "- Avoid using specific job titles; refer to the position as **'this role'** throughout.\n"
            "- Each version should be unique, emphasizing different aspects of the role.\n"
            "- Ensure clarity, conciseness, and engagement in the descriptions.\n\n"
            
            "### Structure for Each Job Description:\n"
            "**1. Role Overview:** A compelling and detailed explanation of this role's significance.\n"
            "**2. Key Responsibilities:** Bullet points outlining core duties, including specifics where applicable.\n"
            "**3. Required Skills:** Essential technical and soft skills, with explanations of their importance.\n"
            "**4. Preferred Skills:** Additional skills that would be advantageous, with context on their relevance.\n"
            "**5. Required Experience:** The necessary experience levels, with examples of relevant past roles.\n"
            "**6. Preferred Experience:** Additional experience that would enhance performance in this role.\n"
            "**7. Tools & Technologies:** Key tools, software, and technologies required for this role.\n"
            "**8. Work Environment & Expectations:** Details on work conditions, methodologies, or collaboration requirements.\n\n"
        
            "Ensure each job description expands on the provided details, enhancing clarity and depth while maintaining industry relevance.\n\n"
            "### Required Format:\n"
            "Present your response exactly as follows:\n\n"
            
            "VERSION 1:\n"
            "[Complete first job description with all sections]\n\n"
            
            "VERSION 2:\n"
            "[Complete second job description with all sections]\n\n"
            
            "VERSION 3:\n"
            "[Complete third job description with all sections]\n\n"
            
            f"### Original Job Description:\n{job_description}\n"
        )

        try:
            native_request = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": self.max_tokens,
                "temperature": self.temperature,
                "messages": [{"role": "user", "content": prompt}],
            }
            
            response = self.client.invoke_model(
                modelId=self.model_id,
                body=json.dumps(native_request),
                contentType="application/json",
            )
            response_body = response['body'].read().decode("utf-8")
            model_response = json.loads(response_body)

            if "content" in model_response and isinstance(model_response["content"], list):
                full_text = model_response["content"][0]["text"].strip()
                
                # More robust splitting pattern
                parts = re.split(r'VERSION \d+:', full_text)
                if len(parts) >= 4:  # The first part is empty or intro text
                    descriptions = [part.strip() for part in parts[1:4]]
                    return descriptions
                else:
                    # Fallback parsing method
                    descriptions = []
                    version_pattern = re.compile(r'VERSION (\d+):(.*?)(?=VERSION \d+:|$)', re.DOTALL)
                    matches = version_pattern.findall(full_text)
                    for _, content in matches[:3]:
                        descriptions.append(content.strip())
                    
                    if len(descriptions) == 3:
                        return descriptions
        except Exception as e:
            print(f"Error generating descriptions: {e}")
        
        # If we failed to parse properly or encountered an error, generate simpler versions
        return [
            f"Enhanced Version 1 of the job description:\n{job_description}",
            f"Enhanced Version 2 of the job description:\n{job_description}",
            f"Enhanced Version 3 of the job description:\n{job_description}"
        ]

    def generate_final_description(self, selected_description, feedback_history):
        """
        Generate enhanced description incorporating feedback history
        
        Args:
            selected_description (str): The base description to enhance
            feedback_history (list): List of previous feedback items
        """
        # If client is not initialized properly, return the selected description
        if not self.client:
            return selected_description + "\n\n[Note: This would normally be enhanced based on your feedback, but the AI service connection is currently unavailable.]"
            
        # Construct prompt with feedback history
        feedback_context = ""
        for i, feedback_item in enumerate(feedback_history[:-1]):
            if isinstance(feedback_item, dict):
                feedback_type = feedback_item.get("type", "General Feedback")
                feedback_text = feedback_item.get("feedback", "")
                feedback_context += f"Previous Feedback {i+1} ({feedback_type}): {feedback_text}\n\n"
            else:
                feedback_context += f"Previous Feedback {i+1}: {feedback_item}\n\n"
        
        # Handle current feedback
        current_feedback = ""
        if feedback_history:
            last_feedback = feedback_history[-1]
            if isinstance(last_feedback, dict):
                feedback_type = last_feedback.get("type", "General Feedback")
                feedback_text = last_feedback.get("feedback", "")
                current_feedback = f"({feedback_type}): {feedback_text}"
            else:
                current_feedback = last_feedback
        
        prompt = (
            "You are an expert in job description refinement. Your task is to enhance the given job description "
            "by incorporating all feedback while maintaining professional quality.\n\n"
            
            f"### Selected Job Description to Enhance:\n{selected_description}\n\n"
        )
        if feedback_context:
            prompt += f"### Previous Feedback Already Incorporated:\n{feedback_context}\n\n"
        
        if current_feedback:
            prompt += f"### New Feedback to Implement:\n{current_feedback}\n\n"
        
        prompt += (
                "### Guidelines:\n"
                "- Implement all feedback while preserving the original core requirements\n"
                "- Maintain clear section structure and professional language\n"
                "- Continue referring to the position as 'this role'\n"
                "- Produce a complete, refined job description ready for immediate use\n\n"
                
                "Return the complete enhanced job description incorporating all feedback."
            )
        
        try:
            native_request = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": self.max_tokens,
                "temperature": self.temperature,
                "messages": [{"role": "user", "content": prompt}],
            }
            
            response = self.client.invoke_model(
                modelId=self.model_id,
                body=json.dumps(native_request),
                contentType="application/json",
            )
            response_body = response['body'].read().decode("utf-8")
            model_response = json.loads(response_body)

            if "content" in model_response and isinstance(model_response["content"], list):
                return model_response["content"][0]["text"].strip()
        except Exception as e:
            print(f"Error generating final description: {e}")
            return selected_description + f"\n\n[Error generating final version: {str(e)}]"

# Resume Analysis and Ranking Functions
def extract_skills(text):
    """Extract technical skills and technologies from text"""
    tech_keywords = {
        'programming_languages': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'scala', 'swift', 'golang'],
        'frameworks': ['django', 'flask', 'spring', 'react', 'angular', 'vue', 'nodejs', 'express', 'hibernate'],
        'databases': ['sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'redis', 'elasticsearch'],
        'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins'],
        'tools': ['git', 'maven', 'gradle', 'junit', 'selenium', 'jira', 'confluence']
    }
    
    text = str(text).lower()
    found_skills = {category: [] for category in tech_keywords}
    
    for category, keywords in tech_keywords.items():
        for keyword in keywords:
            if re.search(r'\b' + keyword + r'\b', text):
                found_skills[category].append(keyword)
    
    return found_skills

def preprocess_text(text):
    """Preprocess text for similarity comparison"""
    if pd.isna(text):
        return ""
    
    text = str(text).lower()
    # Remove special characters and numbers
    text = re.sub(r'[^a-zA-Z\s]', ' ', text)
    # Tokenize
    tokens = word_tokenize(text)
    # Remove stopwords and lemmatize
    tokens = [lemmatizer.lemmatize(token) for token in tokens if token not in stop_words]
    return ' '.join(tokens)

def compute_similarity(job_desc, resume_df):
    """Compute enhanced similarity scores between job description and resumes"""
    # Extract skills from job description
    job_skills = extract_skills(str(job_desc['Skills']) + ' ' + str(job_desc['Tools']))
    
    similarity_scores = []
    for _, resume in resume_df.iterrows():
        # Extract skills from resume
        resume_skills = extract_skills(str(resume['Skills']) + ' ' + str(resume['Tools']))
        
        # Calculate skill matches for each category
        category_scores = []
        for category in job_skills:
            job_set = set(job_skills[category])
            resume_set = set(resume_skills[category])
            if job_set:
                match_ratio = len(resume_set.intersection(job_set)) / len(job_set)
                category_scores.append(match_ratio)
            else:
                category_scores.append(0)
        
        # Calculate weighted skill score
        skill_score = np.mean(category_scores) if category_scores else 0
        
        # Calculate text similarity
        job_text = preprocess_text(str(job_desc['Skills']) + ' ' + str(job_desc['Tools']))
        resume_text = preprocess_text(str(resume['Skills']) + ' ' + str(resume['Tools']) + ' ' + str(resume['Certifications']))
        
        vectorizer = TfidfVectorizer()
        try:
            tfidf_matrix = vectorizer.fit_transform([job_text, resume_text])
            text_similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        except:
            text_similarity = 0
        
        # Combine scores (70% skill match, 30% text similarity)
        final_score = (0.7 * skill_score) + (0.3 * text_similarity)
        similarity_scores.append(final_score)
    
    return np.array(similarity_scores)

def generate_ai_insights(job_desc, resume):
    """Generate AI insights about the resume match using Anthropic or fallback text"""
    if not anthropic_client:
        return f"""
        ### Key Match Analysis:
        
        This candidate shows alignment with the job requirements based on their skills and experience:
        
        - Technical skills match the core requirements
        - Experience with relevant tools and technologies
        - {resume['Certifications'] if resume['Certifications'] else 'Experience'} enhances qualifications
        
        Overall assessment: Good potential match based on technical qualifications.
        """
    
    prompt = f"""
    As an expert in talent evaluation, analyze this job description and resume:
    
    Job Description:
    Skills: {job_desc['Skills']}
    Tools: {job_desc['Tools']}
    
    Resume:
    Skills: {resume['Skills']}
    Tools: {resume['Tools']}
    Certifications: {resume['Certifications']}
    
    Provide:
    1. A concise paragraph on why the candidate is a good fit for this role
    2. 3 key strengths of this candidate's profile (in bullet form)
    3. An overall assessment (Excellent/Good/Fair match)
    """

    try:
        response = anthropic_client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=300,
            temperature=0.7,
            system="You are an expert in HR and talent evaluation.",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response.content[0].text
    except Exception as e:
        print(f"Error generating insights: {e}")
        return f"""
        ### Key Match Analysis:
        
        This candidate shows alignment with the job requirements based on their skills and experience:
        
        - Technical skills match the core requirements
        - Experience with relevant tools and technologies
        - {resume['Certifications'] if resume['Certifications'] else 'Experience'} enhances qualifications
        
        Overall assessment: Good potential match based on technical qualifications.
        
        (Note: This is a fallback analysis as the AI service encountered an error: {str(e)})
        """

def create_radar_chart(resume, job_desc):
    """Create a radar chart for skill matching visualization"""
    categories = ['Technical Skills', 'Tools Proficiency', 'Certifications']
    
    resume_skills = set(str(resume['Skills']).lower().split())
    resume_tools = set(str(resume['Tools']).lower().split())
    resume_certs = set(str(resume['Certifications']).lower().split())
    
    job_skills = set(str(job_desc['Skills']).lower().split())
    job_tools = set(str(job_desc['Tools']).lower().split())
    
    scores = [
        len(resume_skills & job_skills) / max(len(job_skills), 1),
        len(resume_tools & job_tools) / max(len(job_tools), 1),
        len(resume_certs) / 10  # Normalize certification count
    ]
    
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=scores,
        theta=categories,
        fill='toself',
        name='Match Score'
    ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 1]
            )),
        showlegend=False,
        height=250,
        margin=dict(l=20, r=20, t=30, b=20),
        title=None
    )
    return fig

def create_multi_radar_chart(scores_dict):
    """Create a radar chart comparing multiple job descriptions"""
    categories = list(next(iter(scores_dict.values())).keys())
    
    fig = go.Figure()
    
    for label, scores in scores_dict.items():
        fig.add_trace(go.Scatterpolar(
            r=[scores[cat] for cat in categories],
            theta=categories,
            fill='toself',
            name=label
        ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 1]
            )
        ),
        showlegend=True,
        title="Job Description Comparison",
        height=600
    )
    
    return fig

def create_distribution_chart(categorized_resumes):
    """Create a distribution chart showing resume categories"""
    categories = ['High Match', 'Medium Match', 'Low Match']
    counts = [
        len(categorized_resumes['high_matches']),
        len(categorized_resumes['medium_matches']),
        len(categorized_resumes['low_matches'])
    ]
    
    fig = go.Figure(data=[
        go.Bar(
            x=categories,
            y=counts,
            marker_color=['#2ecc71', '#f39c12', '#e74c3c']
        )
    ])
    
    fig.update_layout(
        title="Match Distribution",
        xaxis_title=None,
        yaxis_title="Count",
        height=200,
        margin=dict(l=20, r=20, t=30, b=20),
        showlegend=False
    )
    
    return fig

def create_comparison_dataframe(scores_dict):
    """Create a DataFrame comparing multiple job descriptions"""
    categories = list(next(iter(scores_dict.values())).keys())
    
    df_data = {
        'Category': categories,
    }
    
    # Add scores for each version
    for label, scores in scores_dict.items():
        df_data[label] = [f"{scores[cat]:.2%}" for cat in categories]
        
        # Calculate change from original if this isn't the original
        if label != 'Original':
            original_scores = scores_dict['Original']
            df_data[f'{label} Change'] = [
                f"{(scores[cat] - original_scores[cat])*100:+.2f}%" 
                for cat in categories
            ]
    
    return pd.DataFrame(df_data)

def categorize_resumes(job_desc, resume_df):
    """Categorize resumes into high, medium, and low matches"""
    similarity_scores = compute_similarity(job_desc, resume_df)
    
    all_resumes = []
    for i, score in enumerate(similarity_scores):
        all_resumes.append({
            'Resume ID': resume_df.iloc[i]['File Name'],
            'Skills': resume_df.iloc[i]['Skills'],
            'Tools': resume_df.iloc[i]['Tools'],
            'Certifications': resume_df.iloc[i]['Certifications'],
            'Score': score
        })
    
    # Sort all resumes by score
    all_resumes.sort(key=lambda x: x['Score'], reverse=True)
    
    # Categorize based on score thresholds
    high_matches = [r for r in all_resumes if r['Score'] >= 0.25]
    medium_matches = [r for r in all_resumes if 0.2 <= r['Score'] < 0.25]
    low_matches = [r for r in all_resumes if r['Score'] < 0.2]
    
    return {
        'top_3': all_resumes[:3],
        'high_matches': high_matches,
        'medium_matches': medium_matches,
        'low_matches': low_matches
    }

def init_session_state():
    """Initialize session state variables if they don't exist"""
    if 'role' not in st.session_state:
        st.session_state.role = 'Recruiter'  # Default role
    
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
        
    if 'feedback_history' not in st.session_state:
        st.session_state.feedback_history = []
        
    if 'last_file' not in st.session_state:
        st.session_state.last_file = None
        
    if 'reload_flag' not in st.session_state:
        st.session_state.reload_flag = False
        
    if 'clear_feedback' not in st.session_state:
        st.session_state.clear_feedback = False
        
    if 'viewing_all_feedback' not in st.session_state:
        st.session_state.viewing_all_feedback = False
        
    if 'viewing_session_feedback' not in st.session_state:
        st.session_state.viewing_session_feedback = False
        
    if 'final_version_generated' not in st.session_state:
        st.session_state.final_version_generated = False
        
    if 'final_version' not in st.session_state:
        st.session_state.final_version = None
        
    if 'current_page' not in st.session_state:
        st.session_state.current_page = "jd_enhance"  # Default page
        
    if 'feedback_type' not in st.session_state:
        st.session_state.feedback_type = "General Feedback"
        
    if 'active_tab' not in st.session_state:
        st.session_state.active_tab = "JD Versions"
        
    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = None

def get_or_create_logger():
    """Get existing logger from session state or create a new one"""
    # First check if we have a logger in session state
    if 'logger' in st.session_state:
        return st.session_state.logger
    
    # If we have a session_id, try to load that session
    if 'session_id' in st.session_state:
        try:
            # Try to load existing session by ID
            logger = JDOptimLogger.load_session(st.session_state.session_id)
            if logger:
                # Update role if it changed
                if logger.username != st.session_state.role:
                    logger.username = st.session_state.role
                    logger.current_state["username"] = st.session_state.role
                    logger._save_state()
                
                st.session_state.logger = logger
                return logger
        except Exception as e:
            # If loading fails, we'll create a new logger below
            print(f"Failed to load existing session: {e}")
    
    # Create a new logger with the current role
    logger = JDOptimLogger(username=st.session_state.role)
    st.session_state.session_id = logger.session_id
    st.session_state.logger = logger
    
    return logger

def render_role_selector():
    """Render the role selector in a compact layout"""
    # Display role selector in a small container with border
    with st.container(border=True):
        # Define available roles
        roles = ["Recruiter", "Hiring Manager", "Candidate", "HR Manager", "Team Lead"]
        
        # Simple one-row layout
        selected_role = st.selectbox(
            "Your Role:",
            options=roles,
            index=roles.index(st.session_state.role) if st.session_state.role in roles else 0,
            help="Select your role in the hiring process"
        )
        
        # Update role if changed
        if selected_role != st.session_state.role:
            st.session_state.role = selected_role
            
            # Update logger if it exists
            if 'logger' in st.session_state:
                st.session_state.logger.username = selected_role
                st.session_state.logger.current_state["username"] = selected_role
                st.session_state.logger._save_state()

def display_filtered_feedback_history():
    """Display feedback history with filtering options"""
    # Get all available sessions
    sessions = JDOptimLogger.list_sessions()
    
    if not sessions:
        st.info("No previous feedback found")
        return
    
    # Create a list to store all feedback data
    all_feedback = []
    
    # Collect unique values for filters
    unique_roles = set()
    unique_files = set()
    unique_dates = set()
    unique_feedback_types = set()
    
    # Loop through each session to collect feedback
    for session_info in sessions:
        session_id = session_info["session_id"]
        try:
            # Load the session data directly from file without creating a new logger instance
            log_file = os.path.join("logs", f"jdoptim_session_{session_id}.json")
            if os.path.exists(log_file):
                with open(log_file, 'r') as f:
                    state = json.load(f)
                
                role = state.get("username", "Unknown Role")
                unique_roles.add(role)
                
                file_name = state.get("selected_file", "Unknown")
                unique_files.add(file_name)
                
                session_date = state.get("session_start_time", "Unknown")
                # Extract just the date part if it's a full timestamp
                if isinstance(session_date, str) and "T" in session_date:
                    session_date = session_date.split("T")[0]
                unique_dates.add(session_date)
                
                # Add each feedback item with metadata
                for i, feedback in enumerate(state.get("feedback_history", [])):
                    # Get timestamp for the feedback if available
                    feedback_time = "Unknown"
                    for action in state.get("actions", []):
                        if action.get("action") == "feedback" and action.get("index", -1) == i:
                            feedback_time = action.get("timestamp", "Unknown")
                            break
                    
                    # Handle different feedback formats (string or dict)
                    if isinstance(feedback, dict):
                        feedback_content = feedback.get("feedback", "")
                        feedback_type = feedback.get("type", "General Feedback")
                        feedback_role = feedback.get("role", role)
                    else:
                        feedback_content = feedback
                        feedback_type = "General Feedback"
                        feedback_role = role
                    
                    # Add to unique feedback types
                    unique_feedback_types.add(feedback_type)
                    
                    all_feedback.append({
                        "Role": feedback_role,
                        "File": file_name,
                        "Session Date": session_date,
                        "Feedback Time": feedback_time,
                        "Feedback Type": feedback_type,
                        "Feedback": feedback_content
                    })
        except Exception as e:
            print(f"Error reading session {session_id}: {str(e)}")
    
    if not all_feedback:
        st.info("No feedback found in any session")
        return
            
    # Convert to DataFrame
    feedback_df = pd.DataFrame(all_feedback)
    
    # Sort by most recent first if timestamps are available
    if "Feedback Time" in feedback_df.columns:
        try:
            # Parse timestamps where possible
            parsed_timestamps = []
            for timestamp in feedback_df["Feedback Time"]:
                try:
                    if isinstance(timestamp, str) and "T" in timestamp:
                        dt = datetime.datetime.fromisoformat(timestamp)
                        parsed_timestamps.append(dt)
                    else:
                        parsed_timestamps.append(datetime.datetime(1900, 1, 1))
                except:
                    parsed_timestamps.append(datetime.datetime(1900, 1, 1))
            
            feedback_df["Parsed Timestamp"] = parsed_timestamps
            feedback_df = feedback_df.sort_values("Parsed Timestamp", ascending=False)
        except:
            pass  # If sorting fails, just use the original order
    
    # Create filters with a container to keep UI clean
    with st.expander("Filter Feedback", expanded=False):
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Role filter - make sure we handle None values in the unique_roles set
            cleaned_roles = [role for role in unique_roles if role is not None]
            selected_roles = st.multiselect(
                "Filter by Role:",
                options=sorted(cleaned_roles),
                default=[],
                key="filter_roles"
            )
        
        with col2:
            # File filter - make sure we handle None values in the unique_files set
            cleaned_files = [file for file in unique_files if file is not None]
            selected_files = st.multiselect(
                "Filter by Job Description:",
                options=sorted(cleaned_files),
                default=[],
                key="filter_files"
            )
        
        with col3:
            # Feedback type filter
            cleaned_feedback_types = [ft for ft in unique_feedback_types if ft is not None]
            selected_feedback_types = st.multiselect(
                "Filter by Feedback Type:",
                options=sorted(cleaned_feedback_types),
                default=[],
                key="filter_types"
            )
        
        # Text search
        search_text = st.text_input("Search in feedback:", "", key="search_feedback")
    
    # Apply filters
    filtered_df = feedback_df.copy()
    
    if selected_roles:
        filtered_df = filtered_df[filtered_df["Role"].isin(selected_roles)]
    
    if selected_files:
        filtered_df = filtered_df[filtered_df["File"].isin(selected_files)]
    
    if selected_feedback_types:
        filtered_df = filtered_df[filtered_df["Feedback Type"].isin(selected_feedback_types)]
    
    if search_text:
        filtered_df = filtered_df[filtered_df["Feedback"].str.contains(search_text, case=False, na=False)]
    
    # Show filter summary
    st.write(f"Showing {len(filtered_df)} of {len(feedback_df)} feedback items")
    
    # Display the filtered dataframe
    if not filtered_df.empty:
        # Format timestamps for display
        readable_timestamps = []
        for timestamp in filtered_df["Feedback Time"]:
            try:
                if isinstance(timestamp, str) and "T" in timestamp:
                    dt = datetime.datetime.fromisoformat(timestamp)
                    readable_timestamps.append(dt.strftime("%Y-%m-%d %H:%M:%S"))
                else:
                    readable_timestamps.append(str(timestamp))
            except:
                readable_timestamps.append(str(timestamp))
        
        filtered_df["Formatted Time"] = readable_timestamps
        
        # Display dataframe
        st.dataframe(
            filtered_df,
            use_container_width=True,
            column_config={
                "Role": st.column_config.TextColumn("Role"),
                "File": st.column_config.TextColumn("Job Description"),
                "Formatted Time": st.column_config.TextColumn("Time"),
                "Feedback Type": st.column_config.TextColumn("Feedback Type"),
                "Feedback": st.column_config.TextColumn("Feedback Content", width="large"),
            },
            hide_index=True
        )
    else:
        st.info("No feedback matches the selected filters")
    
    # Option to export filtered results
    if not filtered_df.empty:
        csv = filtered_df.to_csv(index=False)
        st.download_button(
            label="📥 Export Filtered Feedback",
            data=csv,
            file_name=f"feedback_export_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )

def switch_tab(tab_name):
    """Switch between tabs in the application"""
    st.session_state.active_tab = tab_name
    
def switch_page(page_name):
    """Switch between main pages in the application"""
    st.session_state.current_page = page_name

def render_header():
    """Render the application header with logo and title"""
    header_col1, header_col2, header_col3 = st.columns([1, 3, 1])
    
    with header_col1:
        st.markdown(
            """
            <div style="display: flex; justify-content: center; align-items: center; height: 60px;">
                <img src="https://img.icons8.com/color/96/000000/briefcase.png" alt="Dynamic Job Description Optimizer" width="50" height="50">
            </div>
            """, 
            unsafe_allow_html=True
        )
    
    with header_col2:
        st.markdown("<h1 style='text-align: center; margin: 0;'>Dynamic Job Description Optimizer</h1>", unsafe_allow_html=True)
    
    with header_col3:
        st.markdown(
            f"""
            <div style="display: flex; justify-content: center; align-items: center; height: 60px;">
                <div style=padding: 5px 10px; border-radius: 5px; text-align: center;">
                    <div style="font-weight: bold;">{st.session_state.role}</div>
                    <div style="font-size: 0.8em;">View</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

def render_tabs():
    """Render the navigation tabs"""
    tabs = ["JD Versions", "Feedback Loop", "Candidate Ranking", "Client Feedback", "Interview Prep"]

    cols = st.columns(len(tabs))
    for i, tab in enumerate(tabs):
        with cols[i]:
            is_active = st.session_state.active_tab == tab
            if st.button(tab, key=f"tab_{tab}", use_container_width=True, type="secondary" if is_active else "secondary"):
                switch_tab(tab)


def render_jd_enhance_page(logger, analyzer, agent):
    """Render the JD enhancement page"""
    st.markdown(f"<div class='section-header'>📄 Job Description Selection</div>", unsafe_allow_html=True)
    
    jd_directory = os.path.join(os.getcwd(), "JDs")
    try:
        files = [f for f in os.listdir(jd_directory) if f.endswith(('.txt', '.docx'))]
        
        # Create columns for file selection and file preview
        file_col, upload_col = st.columns([2, 1])

        with file_col:
            selected_file = st.selectbox(
                "Select Job Description File",
                files,
                help="Choose a job description file to enhance",
                key="file_selector"
            )

        with upload_col:
            # Add option to upload a file directly
            uploaded_file = st.file_uploader(
                "Or Upload New File",
                type=['txt', 'docx'],
                help="Upload a new job description file"
            )
        
        # Handle file selection or upload
        if uploaded_file:
            if uploaded_file.name.endswith('.txt'):
                content = uploaded_file.getvalue().decode('utf-8')
            else:  # .docx
                # Save to temporary file to use python-docx
                temp_path = f"temp_{uploaded_file.name}"
                with open(temp_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())
                
                # Read content using document function
                doc = Document(temp_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            
            st.session_state.original_jd = content
            selected_file = uploaded_file.name
            
            # Log file selection (only if changed)
            if logger.current_state["selected_file"] != selected_file:
                logger.log_file_selection(selected_file, content)
        elif selected_file:
            file_path = os.path.join(jd_directory, selected_file)
            try:
                # Reset state when file changes
                if st.session_state.last_file != selected_file:
                    st.session_state.last_file = selected_file
                    st.session_state.pop('enhanced_versions', None)
                    st.session_state.pop('original_jd', None)
                    st.session_state.reload_flag = True
                
                # Read the job description
                if 'original_jd' not in st.session_state:
                    st.session_state.original_jd = read_job_description(file_path)
                    
                    # Log file selection (only if changed)
                    if logger.current_state["selected_file"] != selected_file:
                        logger.log_file_selection(selected_file, st.session_state.original_jd)
            except Exception as e:
                st.error(f"Error reading file: {str(e)}")
                return
    except FileNotFoundError:
        # If directory not found, allow direct file upload
        st.warning("Directory 'JDs' not found. You can upload a job description file directly.")
        uploaded_file = st.file_uploader("Upload Job Description File", type=['txt', 'docx'])
        
        if uploaded_file:
            if uploaded_file.name.endswith('.txt'):
                content = uploaded_file.getvalue().decode('utf-8')
            else:  # .docx
                # Save to temporary file to use python-docx
                temp_path = f"temp_{uploaded_file.name}"
                with open(temp_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())
                
                # Read content using document function
                doc = Document(temp_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            
            st.session_state.original_jd = content
            selected_file = uploaded_file.name
            
            # Log file selection (only if changed)
            if logger.current_state["selected_file"] != selected_file:
                logger.log_file_selection(selected_file, content)
        else:
            st.error("Please either upload a file or create a 'JDs' folder in the application directory.")
            return

    # From here, the rest of the app continues with either the uploaded or selected file
    if 'original_jd' in st.session_state:
        original_jd = st.session_state.original_jd
        
        # Display original JD with better formatting
        st.markdown(f"<div class='section-header'>Original Job Description</div>", unsafe_allow_html=True)
        
        # Create tabs for original content and quick preview
        original_tabs = st.tabs(["Full Content", "Quick Preview"])
        
        with original_tabs[0]:
            st.text_area(
                "Original Content", 
                original_jd, 
                height=250, 
                disabled=True,
                key="original_jd_display"
            )
            
        with original_tabs[1]:
            # Show a preview with just first 500 characters
            preview_length = min(500, len(original_jd))
            st.write(original_jd[:preview_length] + ("..." if len(original_jd) > preview_length else ""))
        
        # Generate enhanced versions
        st.markdown(f"<div class='section-header'>✨ Enhanced Versions</div>", unsafe_allow_html=True)
        
        generate_col1, generate_col2 = st.columns([3, 1])
        
        with generate_col1:
            generate_btn = st.button(
                "Generate Enhanced Versions", 
                type="primary", 
                key="generate_btn",
                help="Generate three AI-enhanced versions of your job description"
            )
            
        with generate_col2:
            st.caption("AI will create three distinct improved versions of your job description for you to review.")
        
        # Handle generating enhanced versions
        if generate_btn or ('enhanced_versions' not in st.session_state and st.session_state.reload_flag):
            st.session_state.reload_flag = False
            with st.spinner("Generating enhanced versions... This may take a moment"):
                versions = agent.generate_initial_descriptions(original_jd)
                
                # Ensure we have 3 versions
                while len(versions) < 3:
                    versions.append(f"Enhanced Version {len(versions)+1}:\n{original_jd}")
                
                st.session_state.enhanced_versions = versions
                logger.log_versions_generated(versions)
                st.rerun()

        # Display enhanced versions and their analysis
        if 'enhanced_versions' in st.session_state and len(st.session_state.enhanced_versions) >= 3:
            # Analyze all versions
            original_scores = analyzer.analyze_text(original_jd)
            intermediate_scores = {
                f'Version {i+1}': analyzer.analyze_text(version)
                for i, version in enumerate(st.session_state.enhanced_versions)
            }
            
            # Combine all scores for comparison
            all_scores = {'Original': original_scores, **intermediate_scores}

            # Create tabs for content and analysis
            enhanced_tabs = st.tabs(["Enhanced Versions", "Analysis & Comparison"])
                
            with enhanced_tabs[0]:
                version_tabs = st.tabs(["Version 1", "Version 2", "Version 3"])
                for idx, (tab, version) in enumerate(zip(version_tabs, st.session_state.enhanced_versions)):
                    with tab:
                        st.text_area(
                            f"Enhanced Version {idx + 1}",
                            version,
                            height=300,
                            disabled=True,
                            key=f"enhanced_version_{idx}"
                        )
                        
                        # Add download button for each version
                        st.download_button(
                            label=f"Download Version {idx + 1}",
                            data=version,
                            file_name=f"enhanced_jd_version_{idx+1}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            key=f"download_version_{idx}"
                        )
        
            with enhanced_tabs[1]:
                analysis_col1, analysis_col2 = st.columns([1, 1])
                    
                with analysis_col1:
                    st.subheader("Skill Coverage Comparison")
                    radar_chart = create_multi_radar_chart(all_scores)
                    st.plotly_chart(radar_chart, use_container_width=True, key="intermediate_radar")
                    
                with analysis_col2:
                    st.subheader("Detailed Analysis")
                    comparison_df = create_comparison_dataframe(all_scores)
                    st.dataframe(
                        comparison_df,
                        height=400,
                        use_container_width=True,
                        hide_index=True,
                        key="intermediate_comparison"
                    )
                    st.caption("Percentages indicate keyword coverage in each category")
                    
            # After reviewing enhanced versions, add button to continue to refinement phase
            st.markdown(f"<div class='section-header'>Next Steps</div>", unsafe_allow_html=True)

def render_jd_refine_page(logger, analyzer, agent):
    """Render the JD refinement and feedback page with manual feedback and enhanced outputs"""

    # Breadcrumb for navigation
    breadcrumb_col1, breadcrumb_col2 = st.columns([1, 4])
    with breadcrumb_col1:
        st.markdown(f"<div class='section-header'>🔄 Version Selection & Feedback</div>", unsafe_allow_html=True)

    # Ensure session data exists
    if ('original_jd' not in st.session_state or 
        'enhanced_versions' not in st.session_state or 
        len(st.session_state.enhanced_versions) < 3):
        st.error("Please generate enhanced versions first before proceeding to refinement.")
        if st.button("Go to Generation Page", key="goto_gen"):
            switch_page("jd_enhance")
        return

    # Two-column layout: Left (version selection), Right (feedback input)
    left_col, right_col = st.columns([1, 1])

    with left_col:
        st.markdown(f"<div class='subsection-header'>1. Select Version</div>", unsafe_allow_html=True)
        
        selected_version = st.radio(
            "Choose the version you'd like to use as a base:",
            ["Version 1", "Version 2", "Version 3"],
            help="Select the version that best matches your needs for further enhancement",
            key="version_selector"
        )

        selected_index = int(selected_version[-1]) - 1  # Get version index

        # Display the selected version
        st.text_area(
            f"Selected Version Content",
            st.session_state.enhanced_versions[selected_index],
            height=250,
            disabled=True,
            key=f"selected_version_display"
        )

        # Display previous feedback
        if logger.current_state["feedback_history"]:
            st.markdown(f"<div class='subsection-header'>Previous Feedback</div>", unsafe_allow_html=True)
            with st.expander("View Previous Feedback", expanded=True):
                for i, feedback in enumerate(logger.current_state["feedback_history"], 1):
                    feedback_text = feedback.get("feedback", "") if isinstance(feedback, dict) else feedback
                    feedback_type = feedback.get("type", "General Feedback") if isinstance(feedback, dict) else "General Feedback"
                    st.markdown(f"**#{i} - {feedback_type}:**")
                    st.markdown(f"> {feedback_text}")
                    st.markdown("---")

    with right_col:
        st.markdown(f"<div class='subsection-header'>2. Provide Feedback</div>", unsafe_allow_html=True)

        # Define feedback types
        feedback_types = ["General Feedback", "Rejected Candidate", "Hiring Manager Feedback", 
                          "Client Feedback", "Selected Candidate", "Interview Feedback"]

        selected_feedback_type = st.selectbox(
            "Feedback Type:",
            options=feedback_types,
            index=feedback_types.index(st.session_state.feedback_type) if st.session_state.feedback_type in feedback_types else 0,
            key="feedback_type_selector"
        )

        # Handle manual feedback input
        user_feedback = st.text_area(
            "Enter your suggestions for improving the selected version:",
            height=150,
            placeholder="E.g., 'Add more emphasis on leadership skills', 'Include cloud technologies', etc.",
            key="user_feedback",
            help="Be specific about what you'd like to change or improve"
        )

        # Button to add manual feedback
        if st.button("➕ Add Feedback", type="secondary", key="add_feedback"):
            if user_feedback.strip():
                feedback_obj = {
                    "feedback": user_feedback,
                    "type": selected_feedback_type,
                    "role": st.session_state.role
                }
                logger.current_state["feedback_history"].append(feedback_obj)
                logger._save_state()

                if 'feedback_history' not in st.session_state:
                    st.session_state.feedback_history = []
                st.session_state.feedback_history.append(feedback_obj)

                st.session_state.clear_feedback = True
                st.success("Feedback added successfully!")
                st.rerun()
            else:
                st.warning("Please enter some feedback first.")

    # Generate Final JD Section
    st.markdown(f"<div class='section-header'>🚀 Generate Final Version</div>", unsafe_allow_html=True)

    final_description = ""  # Initialize variable to prevent UnboundLocalError

    if st.button("Generate Final Enhanced JD", type="primary", key="generate_final_jd"):
        try:
            with st.spinner("Enhancing job description with feedback..."):
                logger.log_version_selection(selected_index)

                # Use the current enhanced version if it exists, otherwise use selected version
                base_description = st.session_state.get("current_enhanced_version", 
                                                        st.session_state.enhanced_versions[selected_index])

                # Generate final JD using AI agent
                final_description = agent.generate_final_description(
                    base_description, logger.current_state["feedback_history"]
                )

                # Store new enhanced version in session state
                st.session_state.current_enhanced_version = final_description
                st.session_state.final_version_generated = True
                st.session_state.final_version = final_description

                logger.log_enhanced_version(final_description, is_final=True)
                st.success("Final version generated successfully!")
                st.rerun()
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
    
    # Display Final Version if Generated
    if st.session_state.get('final_version_generated', False) and st.session_state.get('final_version'):
        final_description = st.session_state.final_version  # Retrieve stored JD

        st.markdown("---")
        st.markdown(f"<div class='section-header'>✅ Final Enhanced Job Description</div>", unsafe_allow_html=True)
        st.text_area("Final Content", final_description, height=400, key="final_description")

        # Compare original vs final JD with skill analysis
        original_scores = analyzer.analyze_text(st.session_state.original_jd)
        final_scores = analyzer.analyze_text(final_description)

        st.markdown(f"<div class='section-header'>📊 Final Analysis</div>", unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])
        with col1:
            final_radar = create_multi_radar_chart({'Original': original_scores, 'Final': final_scores})
            st.plotly_chart(final_radar, use_container_width=True, key="final_radar")

        with col2:
            final_comparison_df = create_comparison_dataframe({'Original': original_scores, 'Final': final_scores})
            st.dataframe(final_comparison_df, height=400, use_container_width=True, hide_index=True, key="final_comparison")

        # Download Final JD
        st.markdown(f"<div class='section-header'>📥 Download Options</div>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)

        with col1:
            st.download_button("Download as TXT", data=final_description, file_name=f"enhanced_jd.txt", mime="text/plain", key="download_txt")
            logger.log_download("txt", "enhanced_jd.txt")

        with col2:
            if st.button("Download as DOCX", key="download_docx"):
                save_enhanced_jd(final_description, "enhanced_jd.docx", 'docx')
                st.success("Saved as enhanced_jd.docx")
                logger.log_download("docx", "enhanced_jd.docx")



def render_candidate_ranking_page():
    """Render the candidate ranking page"""
    
    st.markdown(f"<div class='section-header'>🎯 Resume Ranking</div>", unsafe_allow_html=True)
    st.markdown("Raghav please cick on 'Analyze Resumes' to remove the error.", unsafe_allow_html=True)
    # Create sample data if needed
    if not os.path.exists('job_descriptions_analysis_output.csv'):
        st.warning("job_descriptions_analysis_output.csv not found. Using sample data instead.")
        # Create sample job data
        job_data = {
            'File Name': ['DataAnalyticsAIMLJD (1).txt', 'JobDescriptionJavaPythonSupport.txt'],
            'Skills': ['Python, Java, ML, AI, Data Analysis', 'Java, Python, Object-Oriented Programming'],
            'Tools': ['SQL, Cloud, Docker', 'Debugging tools, CoderPad'],
            'JD_Type': ['data_engineer', 'java_developer']
        }
        job_df = pd.DataFrame(job_data)
    else:
        try:
            # Load job data
            job_df = pd.read_csv('job_descriptions_analysis_output.csv')
            
            # Add JD_Type column if it doesn't exist
            if 'JD_Type' not in job_df.columns:
                job_df['JD_Type'] = 'unknown'
                
                # Map Java/Python Support JDs - use more flexible matching
                java_python_keywords = ['java', 'python', 'support']
                
                # Map Data Engineer JDs
                data_engineer_keywords = ['data', 'engineer', 'analytics']
                
                # Apply more flexible mappings
                for index, row in job_df.iterrows():
                    file_name = str(row['File Name']).lower()
                    
                    # Check for Java/Python developer
                    if any(keyword in file_name for keyword in java_python_keywords):
                        job_df.at[index, 'JD_Type'] = 'java_developer'
                    
                    # Check for Data Engineer
                    elif any(keyword in file_name for keyword in data_engineer_keywords):
                        job_df.at[index, 'JD_Type'] = 'data_engineer'
                        
                    # Default mapping
                    else:
                        job_df.at[index, 'JD_Type'] = 'general'
        except Exception as e:
            st.error(f"Error loading job data: {e}")
            # Create sample job data
            job_data = {
                'File Name': ['DataAnalyticsAIMLJD (1).txt', 'JobDescriptionJavaPythonSupport.txt'],
                'Skills': ['Python, Java, ML, AI, Data Analysis', 'Java, Python, Object-Oriented Programming'],
                'Tools': ['SQL, Cloud, Docker', 'Debugging tools, CoderPad'],
                'JD_Type': ['data_engineer', 'java_developer']
            }
            job_df = pd.DataFrame(job_data)
    
    # Create sample resume data if needed
    sample_resume_data = {
        'File Name': ['Resume_1', 'Resume_2', 'Resume_3', 'Resume_4', 'Resume_5'],
        'Skills': [
            'Python, Java, Data Analysis, Machine Learning', 
            'Java, Python, SQL, REST API',
            'C#, .NET, Azure, Cloud Computing',
            'Java, Spring, Hibernate, SQL, REST',
            'Python, ML, AI, Deep Learning, SQL'
        ],
        'Tools': [
            'TensorFlow, Scikit-learn, Docker, Git', 
            'IntelliJ, Eclipse, Git, Maven',
            'Visual Studio, Git, Azure DevOps',
            'Jenkins, Maven, Docker, Kubernetes',
            'Pandas, NumPy, Jupyter, Keras'
        ],
        'Certifications': [
            'AWS Machine Learning Specialty', 
            'Oracle Java Professional',
            'Microsoft Azure Developer',
            'AWS Developer Associate',
            'Google Professional Data Engineer'
        ]
    }
    
    # Create three columns for main layout
    col1, col2, col3 = st.columns([1, 1, 1])

    with col1:
        st.markdown(f"<div class='subsection-header'>Select Position</div>", unsafe_allow_html=True)
        job_desc_file_names = job_df['File Name'].tolist()
        selected_job_desc = st.selectbox('Choose position:', job_desc_file_names, label_visibility="collapsed")
        job_desc = job_df[job_df['File Name'] == selected_job_desc].iloc[0]
        
        # Display the selected JD type for verification
        jd_type = job_desc['JD_Type']
        st.markdown(f"**Resume Pool:** {jd_type.replace('_', ' ').title()}")
        
        with st.expander("Job Details", expanded=False):
            st.markdown(f"**Skills:** {job_desc['Skills']}")
            st.markdown(f"**Tools:** {job_desc['Tools']}")
        
        # Try to load resume data based on the selected job type
        resume_df = None
        possible_resume_files = [
            f'resumes_analysis_output_{jd_type}.csv',
            'resumes_analysis_output.csv',
            os.path.join('Exctracted Resumes', 'resumes_analysis_output.csv'),
            os.path.join('Exctracted Resumes', f'resumes_analysis_output_{jd_type}.csv')
        ]
        
        # Try each possible file path
        for file_path in possible_resume_files:
            try:
                if os.path.exists(file_path):
                    resume_df = pd.read_csv(file_path)
                    st.success(f"Loaded resume data from {file_path}")
                    break
            except Exception as e:
                continue
        
        # If no resume file found, use sample data
        if resume_df is None:
            st.warning("No resume data files found. Using sample data.")
            resume_df = pd.DataFrame(sample_resume_data)
        
        # Analyze button
        if st.button('🔍 Analyze Resumes', type="primary"):
            with st.spinner('Analyzing resumes...'):
                try:
                    categorized_resumes = categorize_resumes(job_desc, resume_df)
                    st.session_state['analysis_results'] = categorized_resumes
                except Exception as e:
                    st.error(f"Error during analysis: {str(e)}")
                    # Create dummy results for demonstration
                    all_resumes = []
                    for i in range(len(resume_df)):
                        score = np.random.uniform(0.1, 0.4)
                        all_resumes.append({
                            'Resume ID': resume_df.iloc[i]['File Name'],
                            'Skills': resume_df.iloc[i]['Skills'],
                            'Tools': resume_df.iloc[i]['Tools'],
                            'Certifications': resume_df.iloc[i]['Certifications'],
                            'Score': score
                        })
                    
                    # Sort by score
                    all_resumes.sort(key=lambda x: x['Score'], reverse=True)
                    
                    # Categorize
                    high_matches = [r for r in all_resumes if r['Score'] >= 0.25]
                    medium_matches = [r for r in all_resumes if 0.2 <= r['Score'] < 0.25]
                    low_matches = [r for r in all_resumes if r['Score'] < 0.2]
                    
                    st.session_state['analysis_results'] = {
                        'top_3': all_resumes[:3],
                        'high_matches': high_matches,
                        'medium_matches': medium_matches,
                        'low_matches': low_matches
                    }

    if 'analysis_results' in st.session_state:
        categorized_resumes = st.session_state['analysis_results']
        
        with col2:
            st.markdown(f"<div class='subsection-header'>Overview</div>", unsafe_allow_html=True)
            # Distribution chart
            try:
                chart = create_distribution_chart(categorized_resumes)
                st.plotly_chart(chart, use_container_width=True)
            except Exception as e:
                st.error(f"Error creating distribution chart: {str(e)}")
                st.bar_chart({
                    'High Match': [len(categorized_resumes['high_matches'])],
                    'Medium Match': [len(categorized_resumes['medium_matches'])],
                    'Low Match': [len(categorized_resumes['low_matches'])]
                })
            
            # Top 3 Quick View
            st.markdown(f"<div class='subsection-header'>Top Matches</div>", unsafe_allow_html=True)
            for i, resume in enumerate(categorized_resumes['top_3'][:3]):
                st.markdown(f"""
                <div class="metric-card">
                    <h4 style="margin:0">#{i + 1} - {resume['Resume ID']}</h4>
                    <p style="margin:0">Match: {resume['Score']:.2%}</p>
                </div>
                """, unsafe_allow_html=True)
        
        with col3:
            st.markdown(f"<div class='subsection-header'>Detailed Analysis</div>", unsafe_allow_html=True)
            tabs = st.tabs(["#1", "#2", "#3"])
            
            for i, (tab, resume) in enumerate(zip(tabs, categorized_resumes['top_3'])):
                with tab:
                    col_a, col_b = st.columns([1, 1])
                    with col_a:
                        st.markdown(f"**Score:** {resume['Score']:.2%}")
                        try:
                            radar_chart = create_radar_chart(resume, job_desc)
                            st.plotly_chart(radar_chart, use_container_width=True)
                        except Exception as e:
                            st.error(f"Error creating radar chart: {str(e)}")
                            st.info("Match analysis visualization unavailable")
                    
                    with col_b:
                        try:
                            insights = generate_ai_insights(job_desc, resume)
                            st.markdown(f"""
                            <div class="insight-box compact-text">
                                {insights}
                            </div>
                            """, unsafe_allow_html=True)
                        except Exception as e:
                            st.error(f"Error generating insights: {str(e)}")
                            st.markdown(f"""
                            <div class="insight-box compact-text">
                                <h4>Key Match Analysis</h4>
                                <p>This candidate has skills that align with the job requirements.</p>
                                <ul>
                                    <li>Technical skills match core requirements</li>
                                    <li>Experience with relevant tools</li>
                                    <li>Professional background enhances qualifications</li>
                                </ul>
                                <p><strong>Overall assessment:</strong> Good potential match</p>
                            </div>
                            """, unsafe_allow_html=True)

        # All Resumes by Category (below the main content)
        st.markdown("---")
        st.markdown(f"<div class='section-header'>📑 All Resumes by Category</div>", unsafe_allow_html=True)
        
        cat_col1, cat_col2, cat_col3 = st.columns(3)
        
        with cat_col1:
            with st.expander(f"High Matches ({len(categorized_resumes['high_matches'])})"):
                for resume in categorized_resumes['high_matches']:
                    st.markdown(f"""
                    <div class="category-high">
                        <h4 style="margin:0">{resume['Resume ID']}</h4>
                        <p style="margin:0">Match: {resume['Score']:.2%}</p>
                    </div>
                    """, unsafe_allow_html=True)
        
        with cat_col2:
            with st.expander(f"Medium Matches ({len(categorized_resumes['medium_matches'])})"):
                for resume in categorized_resumes['medium_matches']:
                    st.markdown(f"""
                    <div class="category-medium">
                        <h4 style="margin:0">{resume['Resume ID']}</h4>
                        <p style="margin:0">Match: {resume['Score']:.2%}</p>
                    </div>
                    """, unsafe_allow_html=True)
        
        with cat_col3:
            with st.expander(f"Low Matches ({len(categorized_resumes['low_matches'])})"):
                for resume in categorized_resumes['low_matches']:
                    st.markdown(f"""
                    <div class="category-low">
                        <h4 style="margin:0">{resume['Resume ID']}</h4>
                        <p style="margin:0">Match: {resume['Score']:.2%}</p>
                    </div>
                    """, unsafe_allow_html=True)

def render_feedback_log_page(logger):
    """Render the feedback log page"""
    st.markdown(f"<div class='section-header'>Feedback Log</div>", unsafe_allow_html=True)
    st.markdown("Track all feedback related to this job", unsafe_allow_html=True)
    
    # Get recent feedback from the current logger and all sessions
    try:
        # Create feedback cards similar to the screenshot
        if logger.current_state["feedback_history"]:
            for i, feedback in enumerate(logger.current_state["feedback_history"], 1):
                feedback_type = "Client"  # Default type for visualization
                feedback_text = ""
                
                if isinstance(feedback, dict):
                    feedback_text = feedback.get("feedback", "")
                    feedback_type = feedback.get("type", "Client")
                    feedback_role = feedback.get("role", logger.username)
                else:
                    feedback_text = feedback
                    feedback_role = logger.username
                
                # Create a card with the appropriate style based on feedback type
                card_color = "#333"
                if "Client" in feedback_type:
                    card_color = "#000"
                elif "Internal" in feedback_type:
                    card_color = "#7e22ce"  # Purple
                elif "Rejected" in feedback_type:
                    card_color = "#e11d48"  # Red
                
                # Format the date consistently
                feedback_date = "Mar 18, 2025"  # Placeholder
                # Try to get actual timestamp from actions
                for action in logger.current_state["actions"]:
                    if action.get("action") == "feedback" and action.get("index", -1) == i-1:
                        try:
                            timestamp = action.get("timestamp", "")
                            if timestamp and isinstance(timestamp, str) and "T" in timestamp:
                                dt = datetime.datetime.fromisoformat(timestamp)
                                feedback_date = dt.strftime("%b %d, %Y")
                        except:
                            pass
                            
                st.markdown(f"""
                <div style="margin-bottom: 20px; border: 1px solid #e5e7eb; border-radius: 8px; overflow: hidden;">
                    <div style="display: flex; align-items: center; padding: 10px; background-color: #f8f9fa;">
                        <div style="background-color: {card_color}; color: white; padding: 5px 10px; border-radius: 5px; margin-right: 10px;">
                            {feedback_type}
                        </div>
                        <div style="flex-grow: 1; font-size: 16px;">
                            {feedback_text[:40] + ('...' if len(feedback_text) > 40 else '')}
                        </div>
                        <div style="color: #6b7280; font-size: 14px;">
                            {feedback_date}
                        </div>
                    </div>
                    <div style="padding: 15px; border-top: 1px solid #e5e7eb;">
                        {feedback_text}
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No feedback available for this job description yet.")
        
        # Add filter and action buttons similar to the screenshot
        col1, col2, col3 = st.columns([1, 1, 2])
        
        with col1:
            st.button("🔍 Filter", key="filter_feedback_btn")
            
        with col2:
            st.button("📤 Export Log", key="export_feedback_btn")
            
        with col3:
            st.button("✏️ Add Feedback", type="primary", key="add_feedback_btn")
            
    except Exception as e:
        st.error(f"Error displaying feedback log: {e}")

def render_interview_prep_page():
    """Render the interview preparation page"""
    st.markdown(f"<div class='section-header'>Interview Preparation</div>", unsafe_allow_html=True)
    
    # Display "Coming Soon" message with a professional look
    st.markdown("""
    <div style="text-align: center; padding: 40px; background-color: #f8f9fa; border-radius: 10px; margin: 20px 0;">
        <img src="https://img.icons8.com/cotton/100/000000/time-machine--v1.png" alt="Coming Soon" width="64" height="64">
        <h2 style="margin-top: 20px; color: #1e3a8a;">Coming Soon</h2>
        <p style="color: #6b7280; max-width: 600px; margin: 0 auto; padding: 10px 0;">
            We're working on an advanced interview preparation module to help you create structured interview 
            questions, evaluation criteria, and candidate scoring templates based on your job descriptions.
        </p>
        <p style="color: #6b7280; max-width: 600px; margin: 10px auto;">
            Stay tuned for updates! This feature will be available in the next release.
        </p>
    </div>
    """, unsafe_allow_html=True)
    

def render_client_feedback_page(agent):
    """Render the Client Feedback tab with JD + Feedback drop zones"""
    st.markdown(f"<div class='section-header'>💬 Client Feedback Enhancement</div>", unsafe_allow_html=True)

    jd_file = st.file_uploader("📄 Drop or upload a Job Description (TXT or DOCX)", type=["txt", "docx"], key="jd_upload")
    feedback_file = st.file_uploader("📝 Drop or upload a Client Feedback File (DOCX)", type=["docx"], key="feedback_upload")

    generate_button = st.button("🚀 Generate Enhanced Job Description", key="generate_client_feedback")

    if generate_button:
        if not jd_file or not feedback_file:
            st.warning("Please upload both the job description and feedback file.")
            return

        # Read job description
        try:
            if jd_file.name.endswith(".txt"):
                job_description = jd_file.read().decode("utf-8")
            elif jd_file.name.endswith(".docx"):
                doc = Document(jd_file)
                job_description = "\n".join([p.text for p in doc.paragraphs])
            else:
                st.error("Unsupported JD file format.")
                return
        except Exception as e:
            st.error(f"Error reading JD file: {str(e)}")
            return

        # Read feedback
        try:
            doc = Document(feedback_file)
            user_feedback = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            st.error(f"Error reading feedback file: {str(e)}")
            return

        # Construct LLM prompt
        prompt = (
            "You are an expert in job description refinement.\n\n"
            "Please revise the provided job description **only based on the feedback** given by the client.\n\n"
            "Do not introduce any information or changes not explicitly stated in the feedback.\n"
            "Only make edits that directly reflect specific feedback content.\n\n"
            "**Guidelines:**\n"
            "- Do not make assumptions.\n"
            "- Do not change formatting or structure unless required by feedback.\n"
            "- Refer to the position as 'this role'.\n"
            "- If the feedback is vague or irrelevant, leave the job description unchanged.\n\n"
            f"### Original Job Description:\n{job_description}\n\n"
            f"### Client Feedback:\n{user_feedback}\n\n"
            "### Please return only the revised job description below (leave unchanged if no edits are needed):\n"
        )

        with st.spinner("Enhancing job description with client feedback..."):
            try:
                native_request = {
                    "anthropic_version": "bedrock-2023-05-31",
                    "max_tokens": 1024,
                    "temperature": 0.7,
                    "messages": [{"role": "user", "content": prompt}],
                }

                response = agent.client.invoke_model(
                    modelId=agent.model_id,
                    body=json.dumps(native_request),
                    contentType="application/json",
                )
                response_body = json.loads(response["body"].read().decode("utf-8"))

                if isinstance(response_body, dict) and "content" in response_body:
                    content = response_body["content"]
                    if isinstance(content, list):
                        enhanced_text = " ".join([item.get("text", "") for item in content])
                    else:
                        enhanced_text = content if isinstance(content, str) else "[No valid content returned]"
                else:
                    enhanced_text = "[Unexpected response format]"

                # Display results
                st.markdown("---")
                st.subheader("📄 Original JD")
                st.text_area("Job Description", job_description, height=200, disabled=True)

                st.subheader("💬 Client Feedback")
                st.text_area("Feedback", user_feedback, height=200, disabled=True)

                st.subheader("✅ Final Enhanced JD")
                st.text_area("Enhanced JD", enhanced_text.strip(), height=400)

            except Exception as e:
                st.error(f"Error generating enhanced JD: {e}")





def main():
    """Main function to run the JD Agent application"""
    st.set_page_config(page_title="JD Agent", page_icon="💼", layout="wide")

    init_session_state()
    logger = get_or_create_logger()
    render_header()
    render_role_selector()
    render_tabs()

    analyzer = JobDescriptionAnalyzer()
    agent = JobDescriptionAgent(model_id="anthropic.claude-3-haiku-20240307-v1:0")

    if st.session_state.active_tab == "JD Versions":
        if st.session_state.current_page == "jd_enhance":
            render_jd_enhance_page(logger, analyzer, agent)
        else:
            render_jd_refine_page(logger, analyzer, agent)
    elif st.session_state.active_tab == "Feedback Loop":
        render_jd_refine_page(logger, analyzer, agent)
    elif st.session_state.active_tab == "Candidate Ranking":
        render_candidate_ranking_page()
    elif st.session_state.active_tab == "Interview Prep":
        render_interview_prep_page()
    elif st.session_state.active_tab == "Client Feedback":
        render_client_feedback_page(agent)

    st.markdown("---")
    st.caption("JD Agent | Made by Apexon")

if __name__ == "__main__":
    main()