import streamlit as st
import boto3
import json
import pandas as pd
import os
from docx import Document
import re
from collections import Counter
import numpy as np
import plotly.graph_objects as go
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import nltk
import datetime
import time
from typing import Dict, List, Any, Optional
import uuid
from jdoptim_logger import JDOptimLogger

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    
def read_job_description(file_path):
    """Read job description from either .txt or .docx file"""
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")

# Function to save enhanced JD
def save_enhanced_jd(content, filename, format_type):
    if format_type == 'docx':
        doc = Document()
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(filename)
        return True
    return False

class JobDescriptionAnalyzer:
    def __init__(self):
        self.categories = {
            'Technical Skills': ['python', 'java', 'sql', 'aws', 'cloud', 'docker', 'kubernetes', 'api', 'database', 
                               'git', 'linux', 'agile', 'devops', 'ml', 'ai', 'analytics'],
            'Soft Skills': ['communication', 'leadership', 'teamwork', 'collaboration', 'problem-solving', 'analytical', 
                           'initiative', 'organizational', 'time management', 'interpersonal'],
            'Experience Level': ['year', 'years', 'senior', 'junior', 'mid-level', 'lead', 'manager', 'experience'],
            'Education': ['degree', 'bachelor', 'master', 'phd', 'certification', 'education'],
            'Tools & Technologies': ['jira', 'confluence', 'slack', 'github', 'gitlab', 'azure', 'jenkins', 'terraform'],
            'Domain Knowledge': ['finance', 'healthcare', 'retail', 'banking', 'insurance', 'technology', 'manufacturing']
        }

    def analyze_text(self, text):
        text = text.lower()
        scores = {}
        
        for category, keywords in self.categories.items():
            category_score = 0
            for keyword in keywords:
                count = len(re.findall(r'\b' + keyword + r'\b', text))
                category_score += count
            max_possible = len(keywords)
            scores[category] = min(category_score / max_possible, 1.0)
            
        return scores

class JobDescriptionAgent:
    def __init__(self, model_id, max_tokens=5000, temperature=0.7):
        self.model_id = model_id
        self.max_tokens = max_tokens
        self.temperature = temperature
        
        # Get AWS credentials from environment variables or use a proper credential provider
        # SECURITY: Replace hardcoded credentials with proper credential management
        self.client = boto3.client(
            service_name='bedrock-runtime',
            aws_access_key_id=st.secrets["aws"]["access_key"],
            aws_secret_access_key=st.secrets["aws"]["secret_key"],
            region_name=st.secrets["aws"]["region"],
        )

    def generate_initial_descriptions(self, job_description):
        """Generate detailed and structured job descriptions based on the given job description."""
        
        prompt = (
            "You are a job description specialist. Your task is to refine and expand upon the provided job description, "
            "creating three distinct versions that are structured, detailed, and aligned with industry best practices.\n\n"
            
            "### Guidelines:\n"
            "- Do NOT make assumptions or introduce inaccuracies.\n"
            "- Avoid using specific job titles; refer to the position as **'this role'** throughout.\n"
            "- Each version should be unique, emphasizing different aspects of the role.\n"
            "- Ensure clarity, conciseness, and engagement in the descriptions.\n\n"
            
            "### Structure for Each Job Description:\n"
            "**1. Role Overview:** A compelling and detailed explanation of this role's significance.\n"
            "**2. Key Responsibilities:** Bullet points outlining core duties, including specifics where applicable.\n"
            "**3. Required Skills:** Essential technical and soft skills, with explanations of their importance.\n"
            "**4. Preferred Skills:** Additional skills that would be advantageous, with context on their relevance.\n"
            "**5. Required Experience:** The necessary experience levels, with examples of relevant past roles.\n"
            "**6. Preferred Experience:** Additional experience that would enhance performance in this role.\n"
            "**7. Tools & Technologies:** Key tools, software, and technologies required for this role.\n"
            "**8. Work Environment & Expectations:** Details on work conditions, methodologies, or collaboration requirements.\n\n"
        
            "Ensure each job description expands on the provided details, enhancing clarity and depth while maintaining industry relevance.\n\n"
            "### Required Format:\n"
            "Present your response exactly as follows:\n\n"
            
            "VERSION 1:\n"
            "[Complete first job description with all sections]\n\n"
            
            "VERSION 2:\n"
            "[Complete second job description with all sections]\n\n"
            
            "VERSION 3:\n"
            "[Complete third job description with all sections]\n\n"
            
            f"### Original Job Description:\n{job_description}\n"
        )


        native_request = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": self.max_tokens,
            "temperature": self.temperature,
            "messages": [{"role": "user", "content": prompt}],
        }
        
        response = self.client.invoke_model(
            modelId=self.model_id,
            body=json.dumps(native_request),
            contentType="application/json",
        )
        response_body = response['body'].read().decode("utf-8")
        model_response = json.loads(response_body)

        if "content" in model_response and isinstance(model_response["content"], list):
            full_text = model_response["content"][0]["text"].strip()
            
            # More robust splitting pattern
            parts = re.split(r'VERSION \d+:', full_text)
            if len(parts) >= 4:  # The first part is empty or intro text
                descriptions = [part.strip() for part in parts[1:4]]
                return descriptions
            else:
                # Fallback parsing method
                descriptions = []
                version_pattern = re.compile(r'VERSION (\d+):(.*?)(?=VERSION \d+:|$)', re.DOTALL)
                matches = version_pattern.findall(full_text)
                for _, content in matches[:3]:
                    descriptions.append(content.strip())
                
                if len(descriptions) == 3:
                    return descriptions
        
        # If we failed to parse properly, generate simpler versions
        return [
            f"Enhanced Version 1 of the job description:\n{job_description}",
            f"Enhanced Version 2 of the job description:\n{job_description}",
            f"Enhanced Version 3 of the job description:\n{job_description}"
        ]

    def generate_final_description(self, selected_description, feedback_history):
        """
        Generate enhanced description incorporating feedback history
        
        Args:
            selected_description (str): The base description to enhance
            feedback_history (list): List of previous feedback items
        """
        # Construct prompt with feedback history
        feedback_context = ""
        for i, feedback_item in enumerate(feedback_history[:-1]):
            if isinstance(feedback_item, dict):
                feedback_type = feedback_item.get("type", "General Feedback")
                feedback_text = feedback_item.get("feedback", "")
                feedback_context += f"Previous Feedback {i+1} ({feedback_type}): {feedback_text}\n\n"
            else:
                feedback_context += f"Previous Feedback {i+1}: {feedback_item}\n\n"
        
        # Handle current feedback
        current_feedback = ""
        if feedback_history:
            last_feedback = feedback_history[-1]
            if isinstance(last_feedback, dict):
                feedback_type = last_feedback.get("type", "General Feedback")
                feedback_text = last_feedback.get("feedback", "")
                current_feedback = f"({feedback_type}): {feedback_text}"
            else:
                current_feedback = last_feedback
        
        prompt = (
            "You are an expert in job description refinement. Your task is to enhance the given job description "
            "by incorporating all feedback while maintaining professional quality.\n\n"
            
            f"### Selected Job Description to Enhance:\n{selected_description}\n\n"
        )
        if feedback_context:
            prompt += f"### Previous Feedback Already Incorporated:\n{feedback_context}\n\n"
        
        if current_feedback:
            prompt += f"### New Feedback to Implement:\n{current_feedback}\n\n"
        
        prompt += (
                "### Guidelines:\n"
                "- Implement all feedback while preserving the original core requirements\n"
                "- Maintain clear section structure and professional language\n"
                "- Continue referring to the position as 'this role'\n"
                "- Produce a complete, refined job description ready for immediate use\n\n"
                
                "Return the complete enhanced job description incorporating all feedback."
            )
        native_request = {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": self.max_tokens,
            "temperature": self.temperature,
            "messages": [{"role": "user", "content": prompt}],
        }
        
        response = self.client.invoke_model(
            modelId=self.model_id,
            body=json.dumps(native_request),
            contentType="application/json",
        )
        response_body = response['body'].read().decode("utf-8")
        model_response = json.loads(response_body)

        if "content" in model_response and isinstance(model_response["content"], list):
            return model_response["content"][0]["text"].strip()
        else:
            raise ValueError("Unexpected response format")

def create_multi_radar_chart(scores_dict):
    """Create a radar chart comparing multiple job descriptions"""
    categories = list(next(iter(scores_dict.values())).keys())
    
    fig = go.Figure()
    
    for label, scores in scores_dict.items():
        fig.add_trace(go.Scatterpolar(
            r=[scores[cat] for cat in categories],
            theta=categories,
            fill='toself',
            name=label
        ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 1]
            )
        ),
        showlegend=True,
        title="Job Description Comparison",
        height=600
    )
    
    return fig

def create_comparison_dataframe(scores_dict):
    """Create a DataFrame comparing multiple job descriptions"""
    categories = list(next(iter(scores_dict.values())).keys())
    
    df_data = {
        'Category': categories,
    }
    
    # Add scores for each version
    for label, scores in scores_dict.items():
        df_data[label] = [f"{scores[cat]:.2%}" for cat in categories]
        
        # Calculate change from original if this isn't the original
        if label != 'Original':
            original_scores = scores_dict['Original']
            df_data[f'{label} Change'] = [
                f"{(scores[cat] - original_scores[cat])*100:+.2f}%" 
                for cat in categories
            ]
    
    return pd.DataFrame(df_data)

def init_session_state():
    """Initialize session state variables if they don't exist"""
    if 'role' not in st.session_state:
        st.session_state.role = 'Recruiter'  # Default role
    
    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
        
    if 'feedback_history' not in st.session_state:
        st.session_state.feedback_history = []
        
    if 'last_file' not in st.session_state:
        st.session_state.last_file = None
        
    if 'reload_flag' not in st.session_state:
        st.session_state.reload_flag = False
        
    if 'clear_feedback' not in st.session_state:
        st.session_state.clear_feedback = False
        
    if 'viewing_all_feedback' not in st.session_state:
        st.session_state.viewing_all_feedback = False
        
    if 'viewing_session_feedback' not in st.session_state:
        st.session_state.viewing_session_feedback = False
        
    if 'final_version_generated' not in st.session_state:
        st.session_state.final_version_generated = False
        
    if 'final_version' not in st.session_state:
        st.session_state.final_version = None
        
    if 'current_page' not in st.session_state:
        st.session_state.current_page = "generate"  # Default page: "generate" or "refine"
        
    if 'feedback_type' not in st.session_state:
        st.session_state.feedback_type = "General Feedback"

def get_or_create_logger():
    """Get existing logger from session state or create a new one"""
    # First check if we have a logger in session state
    if 'logger' in st.session_state:
        return st.session_state.logger
    
    # If we have a session_id, try to load that session
    if 'session_id' in st.session_state:
        try:
            # Try to load existing session by ID
            logger = JDOptimLogger.load_session(st.session_state.session_id)
            if logger:
                # Update role if it changed
                if logger.username != st.session_state.role:
                    logger.username = st.session_state.role
                    logger.current_state["username"] = st.session_state.role
                    logger._save_state()
                
                st.session_state.logger = logger
                return logger
        except Exception as e:
            # If loading fails, we'll create a new logger below
            print(f"Failed to load existing session: {e}")
    
    # Create a new logger with the current role
    logger = JDOptimLogger(username=st.session_state.role)
    st.session_state.session_id = logger.session_id
    st.session_state.logger = logger
    
    return logger

def render_role_selector():
    """Render the role selector in a horizontal layout"""
    st.write("### User Information")
    
    # Define available roles
    roles = ["Recruiter", "Hiring Manager", "Candidate", "HR Manager", "Team Lead"]
    
    # Create columns for role and feedback type
    col1, col2 = st.columns(2)
    
    with col1:
        # Role selection
        selected_role = st.selectbox(
            "Your Role:",
            options=roles,
            index=roles.index(st.session_state.role) if st.session_state.role in roles else 0,
            help="Select your role in the hiring process"
        )
        
        # Update role if changed
        if selected_role != st.session_state.role:
            st.session_state.role = selected_role
            
            # Update logger if it exists
            if 'logger' in st.session_state:
                st.session_state.logger.username = selected_role
                st.session_state.logger.current_state["username"] = selected_role
                st.session_state.logger._save_state()
    
    with col2:
        # Feedback type selection
        feedback_types = ["General Feedback", "Rejected Candidate", "Hiring Manager Feedback", 
                         "Client Feedback", "Selected Candidate", "Interview Feedback"]
        selected_feedback_type = st.selectbox(
            "Default Feedback Type:",
            options=feedback_types,
            index=feedback_types.index(st.session_state.feedback_type) if st.session_state.feedback_type in feedback_types else 0,
            help="Select the default type of feedback you'll be providing"
        )
        
        # Update feedback type if changed
        if selected_feedback_type != st.session_state.feedback_type:
            st.session_state.feedback_type = selected_feedback_type

def display_filtered_feedback_history():
    """Display feedback history with filtering options without creating new sessions"""
    # Get all available sessions
    sessions = JDOptimLogger.list_sessions()
    
    if not sessions:
        st.info("No previous feedback found")
        return
    
    # Create a list to store all feedback data
    all_feedback = []
    
    # Collect unique values for filters
    unique_roles = set()
    unique_files = set()
    unique_dates = set()
    unique_feedback_types = set()
    
    # Loop through each session to collect feedback
    for session_info in sessions:
        session_id = session_info["session_id"]
        try:
            # Load the session data directly from file without creating a new logger instance
            log_file = os.path.join("logs", f"jdoptim_session_{session_id}.json")
            if os.path.exists(log_file):
                with open(log_file, 'r') as f:
                    state = json.load(f)
                
                role = state.get("username", "Unknown Role")
                unique_roles.add(role)
                
                file_name = state.get("selected_file", "Unknown")
                unique_files.add(file_name)
                
                session_date = state.get("session_start_time", "Unknown")
                # Extract just the date part if it's a full timestamp
                if isinstance(session_date, str) and "T" in session_date:
                    session_date = session_date.split("T")[0]
                unique_dates.add(session_date)
                
                # Add each feedback item with metadata
                for i, feedback in enumerate(state.get("feedback_history", [])):
                    # Get timestamp for the feedback if available
                    feedback_time = "Unknown"
                    for action in state.get("actions", []):
                        if action.get("action") == "feedback" and action.get("index", -1) == i:
                            feedback_time = action.get("timestamp", "Unknown")
                            break
                    
                    # Handle different feedback formats (string or dict)
                    if isinstance(feedback, dict):
                        feedback_content = feedback.get("feedback", "")
                        feedback_type = feedback.get("type", "General Feedback")
                        feedback_role = feedback.get("role", role)
                    else:
                        feedback_content = feedback
                        feedback_type = "General Feedback"
                        feedback_role = role
                    
                    # Add to unique feedback types
                    unique_feedback_types.add(feedback_type)
                    
                    all_feedback.append({
                        "Role": feedback_role,
                        "File": file_name,
                        "Session Date": session_date,
                        "Feedback Time": feedback_time,
                        "Feedback Type": feedback_type,
                        "Feedback": feedback_content
                    })
        except Exception as e:
            print(f"Error reading session {session_id}: {str(e)}")
    
    if not all_feedback:
        st.info("No feedback found in any session")
        return
            
    # Convert to DataFrame
    feedback_df = pd.DataFrame(all_feedback)
    
    # Sort by most recent first if timestamps are available
    if "Feedback Time" in feedback_df.columns:
        try:
            # Parse timestamps where possible
            parsed_timestamps = []
            for timestamp in feedback_df["Feedback Time"]:
                try:
                    if isinstance(timestamp, str) and "T" in timestamp:
                        dt = datetime.datetime.fromisoformat(timestamp)
                        parsed_timestamps.append(dt)
                    else:
                        parsed_timestamps.append(datetime.datetime(1900, 1, 1))
                except:
                    parsed_timestamps.append(datetime.datetime(1900, 1, 1))
            
            feedback_df["Parsed Timestamp"] = parsed_timestamps
            feedback_df = feedback_df.sort_values("Parsed Timestamp", ascending=False)
        except:
            pass  # If sorting fails, just use the original order
    
    # Create filters with a container to keep UI clean
    with st.expander("Filter Feedback", expanded=False):
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Role filter - make sure we handle None values in the unique_roles set
            cleaned_roles = [role for role in unique_roles if role is not None]
            selected_roles = st.multiselect(
                "Filter by Role:",
                options=sorted(cleaned_roles),
                default=[],
                key="filter_roles"
            )
        
        with col2:
            # File filter - make sure we handle None values in the unique_files set
            cleaned_files = [file for file in unique_files if file is not None]
            selected_files = st.multiselect(
                "Filter by Job Description:",
                options=sorted(cleaned_files),
                default=[],
                key="filter_files"
            )
        
        with col3:
            # Feedback type filter
            cleaned_feedback_types = [ft for ft in unique_feedback_types if ft is not None]
            selected_feedback_types = st.multiselect(
                "Filter by Feedback Type:",
                options=sorted(cleaned_feedback_types),
                default=[],
                key="filter_types"
            )
        
        # Text search
        search_text = st.text_input("Search in feedback:", "", key="search_feedback")
    
    # Apply filters
    filtered_df = feedback_df.copy()
    
    if selected_roles:
        filtered_df = filtered_df[filtered_df["Role"].isin(selected_roles)]
    
    if selected_files:
        filtered_df = filtered_df[filtered_df["File"].isin(selected_files)]
    
    if selected_feedback_types:
        filtered_df = filtered_df[filtered_df["Feedback Type"].isin(selected_feedback_types)]
    
    if search_text:
        filtered_df = filtered_df[filtered_df["Feedback"].str.contains(search_text, case=False, na=False)]
    
    # Show filter summary
    st.write(f"Showing {len(filtered_df)} of {len(feedback_df)} feedback items")
    
    # Display the filtered dataframe
    if not filtered_df.empty:
        # Format timestamps for display
        readable_timestamps = []
        for timestamp in filtered_df["Feedback Time"]:
            try:
                if isinstance(timestamp, str) and "T" in timestamp:
                    dt = datetime.datetime.fromisoformat(timestamp)
                    readable_timestamps.append(dt.strftime("%Y-%m-%d %H:%M:%S"))
                else:
                    readable_timestamps.append(str(timestamp))
            except:
                readable_timestamps.append(str(timestamp))
        
        filtered_df["Formatted Time"] = readable_timestamps
        
        # Display dataframe
        st.dataframe(
            filtered_df,
            use_container_width=True,
            column_config={
                "Role": st.column_config.TextColumn("Role"),
                "File": st.column_config.TextColumn("Job Description"),
                "Formatted Time": st.column_config.TextColumn("Time"),
                "Feedback Type": st.column_config.TextColumn("Feedback Type"),
                "Feedback": st.column_config.TextColumn("Feedback Content", width="large"),
            },
            hide_index=True
        )
    else:
        st.info("No feedback matches the selected filters")
    
    # Option to export filtered results
    if not filtered_df.empty:
        csv = filtered_df.to_csv(index=False)
        st.download_button(
            label="📥 Export Filtered Feedback",
            data=csv,
            file_name=f"feedback_export_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )

def cleanup_old_logs(max_sessions=50, older_than_days=30):
    """Clean up old log files"""
    log_dir = "logs"
    if not os.path.exists(log_dir):
        return
    
    # Get all log files with their modification times
    files_with_time = []
    for filename in os.listdir(log_dir):
        if filename.startswith("jdoptim_session_") and filename.endswith(".json"):
            file_path = os.path.join(log_dir, filename)
            mod_time = os.path.getmtime(file_path)
            files_with_time.append((file_path, mod_time))
    
    # Sort by modification time (newest first)
    files_with_time.sort(key=lambda x: x[1], reverse=True)
    
    # Keep recent sessions (max_sessions)
    keep_files = [file_path for file_path, _ in files_with_time[:max_sessions]]
    
    # Delete older files
    current_time = time.time()
    cutoff_time = current_time - (older_than_days * 24 * 60 * 60)
    
    deleted_count = 0
    for file_path, mod_time in files_with_time:
        # Only delete if it's not in the keep_files list and older than cutoff
        if file_path not in keep_files and mod_time < cutoff_time:
            try:
                os.remove(file_path)
                deleted_count += 1
            except:
                pass
    
    return deleted_count

def start_new_session():
    """Start a new session"""
    # Clean up session state
    for key in ['logger', 'session_id', 'enhanced_versions', 'original_jd', 
                'feedback_history', 'current_enhanced_version', 'last_file']:
        if key in st.session_state:
            del st.session_state[key]
    
    # Also reset current page to generate
    st.session_state.current_page = "generate"
    
    # Flag for reload
    st.session_state.reload_flag = True

def switch_to_refinement_page():
    """Switch to the refinement page"""
    st.session_state.current_page = "refine"
    st.rerun()

def switch_to_generation_page():
    """Switch to the generation page"""
    st.session_state.current_page = "generate"
    st.rerun()

def display_help_section():
    """Display a simple horizontal workflow help text in an expander"""
    with st.expander("📚 How to Use This Tool", expanded=False):
        # Use a simple horizontal text layout
        help_html = """
        <style>
        .horizontal-workflow {
            width: 100%;
            white-space: nowrap;
            overflow-x: auto;
            padding: 10px;
            font-size: 14px;
            line-height: 1.5;
        }
        .workflow-title {
            font-weight: bold;
            margin-bottom: 10px;
        }
        .workflow-steps {
            display: flex;
            align-items: center;
        }
        .step {
            margin-right: 5px;
        }
        .arrow {
            margin: 0 5px;
            color: #666;
        }
        .tips {
            margin-top: 10px;
            font-style: italic;
        }
        </style>
        
        <div class="horizontal-workflow">
            <div class="workflow-title">JD Enhancer Workflow:</div>
            <div class="workflow-steps">
                <span class="step">1. Select Role</span>
                <span class="arrow">→</span>
                <span class="step">2. Choose File</span>
                <span class="arrow">→</span>
                <span class="step">3. Generate Versions</span>
                <span class="arrow">→</span>
                <span class="step">4. Review</span>
                <span class="arrow">→</span>
                <span class="step">5. Select Version</span>
                <span class="arrow">→</span>
                <span class="step">6. Provide Feedback</span>
                <span class="arrow">→</span>
                <span class="step">7. Generate Final</span>
                <span class="arrow">→</span>
                <span class="step">8. Download</span>
            </div>
            <div class="tips">
                Tips: Be specific in feedback • Consider different aspects • Use analysis charts for balanced coverage
            </div>
        </div>
        """
        
        # Display the simple horizontal workflow
        st.markdown(help_html, unsafe_allow_html=True)

def render_generation_page(logger, analyzer, agent):
    """Render the JD generation page"""
    st.markdown("### 📄 Job Description Selection")
    
    jd_directory = os.path.join(os.getcwd(), "JDs")
    try:
        files = [f for f in os.listdir(jd_directory) if f.endswith(('.txt', '.docx'))]
        
        # Create columns for file selection and file preview
        file_col, upload_col = st.columns([2, 1])

        with file_col:
            selected_file = st.selectbox(
                "Select Job Description File",
                files,
                help="Choose a job description file to enhance",
                key="file_selector"
            )

        with upload_col:
            # Add option to upload a file directly
            uploaded_file = st.file_uploader(
                "Or Upload New File",
                type=['txt', 'docx'],
                help="Upload a new job description file"
            )
        
        # Handle file selection or upload
        if uploaded_file:
            if uploaded_file.name.endswith('.txt'):
                content = uploaded_file.getvalue().decode('utf-8')
            else:  # .docx
                # Save to temporary file to use python-docx
                temp_path = f"temp_{uploaded_file.name}"
                with open(temp_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())
                
                # Read content using document function
                doc = Document(temp_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            
            st.session_state.original_jd = content
            selected_file = uploaded_file.name
            
            # Log file selection (only if changed)
            if logger.current_state["selected_file"] != selected_file:
                logger.log_file_selection(selected_file, content)
        elif selected_file:
            file_path = os.path.join(jd_directory, selected_file)
            try:
                # Reset state when file changes
                if st.session_state.last_file != selected_file:
                    st.session_state.last_file = selected_file
                    st.session_state.pop('enhanced_versions', None)
                    st.session_state.pop('original_jd', None)
                    st.session_state.reload_flag = True
                
                # Read the job description
                if 'original_jd' not in st.session_state:
                    st.session_state.original_jd = read_job_description(file_path)
                    
                    # Log file selection (only if changed)
                    if logger.current_state["selected_file"] != selected_file:
                        logger.log_file_selection(selected_file, st.session_state.original_jd)
            except Exception as e:
                st.error(f"Error reading file: {str(e)}")
                return
    except FileNotFoundError:
        # If directory not found, allow direct file upload
        st.warning("Directory 'JDs' not found. You can upload a job description file directly.")
        uploaded_file = st.file_uploader("Upload Job Description File", type=['txt', 'docx'])
        
        if uploaded_file:
            if uploaded_file.name.endswith('.txt'):
                content = uploaded_file.getvalue().decode('utf-8')
            else:  # .docx
                # Save to temporary file to use python-docx
                temp_path = f"temp_{uploaded_file.name}"
                with open(temp_path, 'wb') as f:
                    f.write(uploaded_file.getvalue())
                
                # Read content using document function
                doc = Document(temp_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            
            st.session_state.original_jd = content
            selected_file = uploaded_file.name
            
            # Log file selection (only if changed)
            if logger.current_state["selected_file"] != selected_file:
                logger.log_file_selection(selected_file, content)
        else:
            st.error("Please either upload a file or create a 'JDs' folder in the application directory.")
            return

    # From here, the rest of the app continues with either the uploaded or selected file
    if 'original_jd' in st.session_state:
        original_jd = st.session_state.original_jd
        
        # Display original JD with better formatting
        st.markdown("### Original Job Description")
        
        # Create tabs for original content and quick preview
        original_tabs = st.tabs(["Full Content", "Quick Preview"])
        
        with original_tabs[0]:
            st.text_area(
                "Original Content", 
                original_jd, 
                height=250, 
                disabled=True,
                key="original_jd_display"
            )
            
        with original_tabs[1]:
            # Show a preview with just first 500 characters
            preview_length = min(500, len(original_jd))
            st.write(original_jd[:preview_length] + ("..." if len(original_jd) > preview_length else ""))
        
        # Generate enhanced versions
        st.markdown("### ✨ Enhanced Versions")
        
        generate_col1, generate_col2 = st.columns([3, 1])
        
        with generate_col1:
            generate_btn = st.button(
                "Generate Enhanced Versions", 
                type="primary", 
                key="generate_btn",
                help="Generate three AI-enhanced versions of your job description"
            )
            
        with generate_col2:
            st.caption("AI will create three distinct improved versions of your job description for you to review.")
        
        # Handle generating enhanced versions
        if generate_btn or ('enhanced_versions' not in st.session_state and st.session_state.reload_flag):
            st.session_state.reload_flag = False
            with st.spinner("Generating enhanced versions... This may take a moment"):
                versions = agent.generate_initial_descriptions(original_jd)
                
                # Ensure we have 3 versions
                while len(versions) < 3:
                    versions.append(f"Enhanced Version {len(versions)+1}:\n{original_jd}")
                
                st.session_state.enhanced_versions = versions
                logger.log_versions_generated(versions)
                st.rerun()

        # Display enhanced versions and their analysis
        if 'enhanced_versions' in st.session_state and len(st.session_state.enhanced_versions) >= 3:
            # Analyze all versions
            original_scores = analyzer.analyze_text(original_jd)
            intermediate_scores = {
                f'Version {i+1}': analyzer.analyze_text(version)
                for i, version in enumerate(st.session_state.enhanced_versions)
            }
            
            # Combine all scores for comparison
            all_scores = {'Original': original_scores, **intermediate_scores}

            # Create tabs for content and analysis
            enhanced_tabs = st.tabs(["Enhanced Versions", "Analysis & Comparison"])
                
            with enhanced_tabs[0]:
                version_tabs = st.tabs(["Version 1", "Version 2", "Version 3"])
                for idx, (tab, version) in enumerate(zip(version_tabs, st.session_state.enhanced_versions)):
                    with tab:
                        st.text_area(
                            f"Enhanced Version {idx + 1}",
                            version,
                            height=300,
                            disabled=True,
                            key=f"enhanced_version_{idx}"
                        )
                        
                        # Add download button for each version
                        st.download_button(
                            label=f"Download Version {idx + 1}",
                            data=version,
                            file_name=f"enhanced_jd_version_{idx+1}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            key=f"download_version_{idx}"
                        )
        
            with enhanced_tabs[1]:
                analysis_col1, analysis_col2 = st.columns([1, 1])
                    
                with analysis_col1:
                    st.subheader("Skill Coverage Comparison")
                    radar_chart = create_multi_radar_chart(all_scores)
                    st.plotly_chart(radar_chart, use_container_width=True, key="intermediate_radar")
                    
                with analysis_col2:
                    st.subheader("Detailed Analysis")
                    comparison_df = create_comparison_dataframe(all_scores)
                    st.dataframe(
                        comparison_df,
                        height=400,
                        use_container_width=True,
                        hide_index=True,
                        key="intermediate_comparison"
                    )
                    st.caption("Percentages indicate keyword coverage in each category")
                    
            # After reviewing enhanced versions, add button to continue to refinement phase
            st.markdown("### Next Steps")
            
            refinement_col1, refinement_col2 = st.columns(2)
            
            with refinement_col1:
                if st.button("Continue to Version Selection & Feedback", type="primary"):
                    switch_to_refinement_page()
            
            with refinement_col2:
                st.caption("Proceed to select your preferred version and provide feedback to further refine the job description.")

def render_refinement_page(logger, analyzer, agent):
    """Render the JD refinement and feedback page"""
    
    # Add navigation breadcrumb with return option
    breadcrumb_col1, breadcrumb_col2 = st.columns([1, 4])
    with breadcrumb_col1:
        if st.button("← Back to Version Generation", key="back_btn"):
            switch_to_generation_page()
    
    with breadcrumb_col2:
        st.markdown("### 🔄 Version Selection & Feedback")
    
    # Make sure we have the session data we need
    if ('original_jd' not in st.session_state or 
        'enhanced_versions' not in st.session_state or 
        len(st.session_state.enhanced_versions) < 3):
        st.error("Please generate enhanced versions first before proceeding to refinement.")
        if st.button("Go to Generation Page", key="goto_gen"):
            switch_to_generation_page()
        return
    
    # Setup the layout with two columns
    left_col, right_col = st.columns([1, 1])
    
    # In the left column, show the version selection and previous feedback
    with left_col:
        st.subheader("1. Select Version")
        selected_version = st.radio(
            "Choose the version you'd like to use as a base:",
            ["Version 1", "Version 2", "Version 3"],
            help="Select the version that best matches your needs for further enhancement",
            key="version_selector"
        )
        
        # Get selected version index
        selected_index = int(selected_version[-1]) - 1
        
        # Display the selected version
        st.text_area(
            f"Selected Version Content",
            st.session_state.enhanced_versions[selected_index],
            height=250,
            disabled=True,
            key=f"selected_version_display"
        )
        
        # Display previous feedback if available
        if logger.current_state["feedback_history"]:
            st.subheader("Previous Feedback")
            with st.expander("View Previous Feedback", expanded=True):
                for i, feedback in enumerate(logger.current_state["feedback_history"], 1):
                    if isinstance(feedback, dict):
                        feedback_text = feedback.get("feedback", "")
                        feedback_type = feedback.get("type", "General Feedback")
                        st.markdown(f"**#{i} - {feedback_type}:**")
                        st.markdown(f"> {feedback_text}")
                        st.markdown("---")
                    else:
                        st.markdown(f"**#{i}:**")
                        st.markdown(f"> {feedback}")
                        st.markdown("---")
    
    # In the right column, show the feedback input and submission options
    with right_col:
        st.subheader("2. Provide Feedback")
        
        # Define feedback type options
        feedback_types = ["General Feedback", "Rejected Candidate", "Hiring Manager Feedback", 
                         "Client Feedback", "Selected Candidate", "Interview Feedback"]
        
        # Select feedback type
        selected_feedback_type = st.selectbox(
            "Feedback Type:",
            options=feedback_types,
            index=feedback_types.index(st.session_state.feedback_type) if st.session_state.feedback_type in feedback_types else 0,
            key="feedback_type_selector"
        )
        
        # Handle feedback clearing mechanism
        if st.session_state.get('clear_feedback', False):
            st.session_state.clear_feedback = False
        
        # Input for manual feedback
        user_feedback = st.text_area(
            "Enter your suggestions for improving the selected version:",
            height=150,
            placeholder="E.g., 'Add more emphasis on leadership skills', 'Include cloud technologies', 'Remove references to specific programming languages', etc.",
            key="user_feedback",
            help="Be specific about what you'd like to change or improve"
        )
        
        # Add divider before feedback file option
        st.markdown("--- OR ---")
        
        # Option to select feedback from file
        st.write("Select feedback from a file:")
        
        # Create tabs for selecting from directory or uploading
        feedback_tabs = st.tabs(["Select from Feedbacks Folder", "Upload Feedback File"])
        
        feedback_from_file = None
        file_feedback_type = selected_feedback_type
        
        with feedback_tabs[0]:
            # Select from Feedbacks folder
            feedback_directory = os.path.join(os.getcwd(), "Feedbacks")
            
            # Check if directory exists
            if not os.path.exists(feedback_directory):
                st.warning("The 'Feedbacks' directory does not exist. Please create it or upload a file directly.")
            else:
                # Get all .txt and .docx files from the Feedbacks folder
                feedback_files = [f for f in os.listdir(feedback_directory) 
                               if f.endswith(('.txt', '.docx'))]
                
                if not feedback_files:
                    st.warning("No feedback files found in the Feedbacks directory. Please add .txt or .docx files.")
                else:
                    # Allow user to select a feedback file
                    selected_feedback_file = st.selectbox(
                        "Select Feedback File",
                        feedback_files,
                        help="Choose a feedback file to process"
                    )
                    
                    # Select feedback type for file
                    file_feedback_type = st.selectbox(
                        "File Feedback Type:",
                        options=feedback_types,
                        index=feedback_types.index(selected_feedback_type),
                        key="file_feedback_type"
                    )
                    
                    if selected_feedback_file:
                        feedback_path = os.path.join(feedback_directory, selected_feedback_file)
                        
                        # Extract text based on file type
                        try:
                            if selected_feedback_file.endswith('.txt'):
                                with open(feedback_path, 'r', encoding='utf-8') as file:
                                    feedback_from_file = file.read()
                            elif selected_feedback_file.endswith('.docx'):
                                doc = Document(feedback_path)
                                feedback_from_file = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                            
                            # Display the feedback content
                            st.text_area(
                                "Feedback Content",
                                feedback_from_file,
                                height=100,
                                disabled=True
                            )
                        except Exception as e:
                            st.error(f"Error reading feedback file: {str(e)}")
        
        with feedback_tabs[1]:
            # Upload file option
            uploaded_feedback = st.file_uploader(
                "Upload Feedback File",
                type=['txt', 'docx'],
                help="Upload a .txt or .docx file containing feedback"
            )
            
            # Select feedback type for uploaded file
            upload_feedback_type = st.selectbox(
                "Uploaded File Feedback Type:",
                options=feedback_types,
                index=feedback_types.index(selected_feedback_type),
                key="upload_feedback_type"
            )
            
            if uploaded_feedback:
                try:
                    # Extract text based on file type
                    if uploaded_feedback.name.endswith('.txt'):
                        feedback_from_file = uploaded_feedback.getvalue().decode('utf-8')
                    elif uploaded_feedback.name.endswith('.docx'):
                        # Save to temporary file to use python-docx
                        temp_path = f"temp_{uploaded_feedback.name}"
                        with open(temp_path, 'wb') as f:
                            f.write(uploaded_feedback.getvalue())
                        
                        doc = Document(temp_path)
                        feedback_from_file = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                        
                        # Clean up temp file
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    
                    # Use the feedback type from the upload section
                    file_feedback_type = upload_feedback_type
                    
                    # Display the feedback content
                    st.text_area(
                        "Feedback Content",
                        feedback_from_file,
                        height=100,
                        disabled=True
                    )
                except Exception as e:
                    st.error(f"Error processing uploaded file: {str(e)}")
        
        # Create a row for feedback action buttons
        feedback_btn_col1, feedback_btn_col2, feedback_btn_col3 = st.columns(3)
        
        with feedback_btn_col1:
            # Add a button to save manual feedback without generating final version
            if st.button("Add Manual Feedback", type="secondary", key="add_feedback_only"):
                if user_feedback.strip():
                    # Create feedback object with type
                    feedback_obj = {
                        "feedback": user_feedback,
                        "type": selected_feedback_type,
                        "role": st.session_state.role
                    }
                    
                    # Add to the logger's feedback history directly
                    logger.current_state["feedback_history"].append(feedback_obj)
                    
                    # Save the updated state
                    logger._save_state()
                    
                    # Add to session state for UI update
                    if 'feedback_history' not in st.session_state:
                        st.session_state.feedback_history = []
                    st.session_state.feedback_history.append(feedback_obj)
                    
                    # Log the action
                    logger.current_state["actions"].append({
                        "action": "feedback",
                        "timestamp": datetime.datetime.now().isoformat(),
                        "index": len(logger.current_state["feedback_history"]) - 1
                    })
                    logger._save_state()
                    
                    # Clear feedback input
                    st.session_state.clear_feedback = True
                    st.success("Feedback added successfully! You can add more feedback or generate the final version when ready.")
                    st.rerun()
                else:
                    st.warning("Please enter some feedback first.")
        
        with feedback_btn_col2:
            # Add a button to use feedback from file
            if st.button("Add File Feedback", type="secondary", key="add_file_feedback"):
                if feedback_from_file:
                    # Create feedback object with type
                    feedback_obj = {
                        "feedback": feedback_from_file,
                        "type": file_feedback_type,
                        "role": st.session_state.role
                    }
                    
                    # Add to the logger's feedback history directly
                    logger.current_state["feedback_history"].append(feedback_obj)
                    
                    # Save the updated state
                    logger._save_state()
                    
                    # Add to session state for UI update
                    if 'feedback_history' not in st.session_state:
                        st.session_state.feedback_history = []
                    st.session_state.feedback_history.append(feedback_obj)
                    
                    # Log the action
                    logger.current_state["actions"].append({
                        "action": "feedback",
                        "timestamp": datetime.datetime.now().isoformat(),
                        "index": len(logger.current_state["feedback_history"]) - 1
                    })
                    logger._save_state()
                    
                    st.success("File feedback added successfully!")
                    st.rerun()
                else:
                    st.warning("Please select or upload a feedback file first.")
        
        with feedback_btn_col3:
            # View all feedback button
            if st.button("View All Feedback", type="secondary", key="view_all_feedback"):
                st.session_state['viewing_all_feedback'] = True
    
    # Display all feedback if requested
    if st.session_state.get('viewing_all_feedback', False):
        st.markdown("### 📋 All Feedback History")
        display_filtered_feedback_history()
        # Reset viewing flag after displaying
        st.session_state['viewing_all_feedback'] = False
    
    # Final enhancement process section
    st.markdown("### 🚀 Generate Final Version")
    
    final_col1, final_col2 = st.columns(2)
        
    with final_col1:
        current_feedback_count = len(logger.current_state['feedback_history'])
        current_btn_label = f"Generate Final Version ({current_feedback_count} feedback items)"
        if st.button(current_btn_label, type="primary", key="generate_current_feedback"):
            try:
                with st.spinner("Creating final version with feedback... This may take a moment"):
                    # Log version selection
                    logger.log_version_selection(selected_index)
                        
                    # Log new feedback if provided and not already added
                    if user_feedback.strip():
                        # Create feedback object with type
                        feedback_obj = {
                            "feedback": user_feedback,
                            "type": selected_feedback_type,
                            "role": st.session_state.role
                        }
                        
                        # Add to the logger's feedback history directly
                        logger.current_state["feedback_history"].append(feedback_obj)
                        
                        # Save the updated state
                        logger._save_state()
                        
                        # Log the action
                        logger.current_state["actions"].append({
                            "action": "feedback",
                            "timestamp": datetime.datetime.now().isoformat(),
                            "index": len(logger.current_state["feedback_history"]) - 1
                        })
                        logger._save_state()
                        
                        # Clear feedback input
                        st.session_state.clear_feedback = True
                        
                    # Use the current enhanced version if it exists, otherwise use selected version
                    base_description = (logger.current_state["current_enhanced_version"] or 
                                        st.session_state.enhanced_versions[selected_index])
                        
                    # Generate final description using current session feedback
                    final_description = agent.generate_final_description(
                        base_description,
                        logger.current_state["feedback_history"]
                    )
                        
                    # Log the new enhanced version
                    logger.log_enhanced_version(final_description, is_final=True)
                        
                    # Store the new enhanced version in session state
                    st.session_state.current_enhanced_version = final_description
                        
                    # Flag that we generated the final version
                    st.session_state.final_version_generated = True
                    st.session_state.final_version = final_description
                        
                    st.success("Final version generated successfully!")
                    st.rerun()
                        
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
                st.error("Please try again or contact support if the problem persists.")
        
    with final_col2:
        st.caption("""
        Generate the final enhanced version based on your selected template and all feedback provided. 
        The AI will incorporate all feedback items to create an optimized job description that meets your requirements.
        """)
        
    # Display final version if it was generated
    if st.session_state.get('final_version_generated', False) and st.session_state.get('final_version'):
        final_description = st.session_state.final_version
            
        # Display final version with clear separation
        st.markdown("---")
        st.markdown("### ✅ Final Enhanced Job Description")
        
        # Create a container with a border for the final version
        final_container = st.container(border=True)
        with final_container:
            st.text_area(
                "Final Content",
                final_description,
                height=400,
                key="final_description"
            )
            
        # Add final version to scores dictionary for comparison
        original_scores = analyzer.analyze_text(st.session_state.original_jd)
        final_scores = analyzer.analyze_text(final_description)
        
        # Create comparison section
        st.markdown("### 📊 Final Analysis")
        
        final_col1, final_col2 = st.columns([1, 1])
            
        with final_col1:
            final_radar = create_multi_radar_chart({'Original': original_scores, 'Final': final_scores})
            st.plotly_chart(final_radar, use_container_width=True, key="final_radar")
            
        with final_col2:
            # Create a simplified comparison for final vs original
            final_comparison_df = create_comparison_dataframe({'Original': original_scores, 'Final': final_scores})
            st.dataframe(
                final_comparison_df,
                height=400,
                use_container_width=True,
                hide_index=True,
                key="final_comparison"
            )
            
        # Download section
        st.markdown("### 📥 Download Options")
        
        download_col1, download_col2 = st.columns(2)
            
        with download_col1:
            st.download_button(
                label="Download as TXT",
                data=final_description,
                file_name=f"enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                key="download_txt"
            )
            # Log download action
            logger.log_download("txt", f"enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
            
        with download_col2:
            if st.button("Download as DOCX", key="download_docx"):
                docx_filename = f"enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                save_enhanced_jd(final_description, docx_filename, 'docx')
                st.success(f"Saved as {docx_filename}")
                # Log download action
                logger.log_download("docx", docx_filename)

def main():
    # Use a clean Streamlit theme
    st.set_page_config(
        page_title="Job Description Enhancer",
        page_icon="💼",
        layout="wide",
        initial_sidebar_state="collapsed"
    )

    # Initialize session state
    init_session_state()
    
    # Get or create logger
    logger = get_or_create_logger()
    
    # Clean up old logs in the background
    cleanup_old_logs()

    # Display header with title and role selector in same row
    st.markdown("<h1 style='text-align: center;'>💼 Job Description Enhancer</h1>", unsafe_allow_html=True)
    
    # Display help button in the top right
    help_col, title_col = st.columns([1, 5])
    with help_col:
        display_help_section()
    
    # Space before role selector
    st.write("")

    # Render role selector
    render_role_selector()
    
    # Session management in a cleaner horizontal layout
    session_col1, session_col2, session_col3 = st.columns(3)
    
    with session_col1:
        st.button("🔄 Start New Session", on_click=start_new_session)
        
    with session_col2:
        st.caption(f"Current Session ID: {logger.session_id[:8]}...")
        st.caption(f"Role: {logger.username}")
        
    with session_col3:
        if logger.current_state["actions"]:
            st.caption(f"Actions in this session: {len(logger.current_state['actions'])}")
            if logger.current_state["selected_file"]:
                st.caption(f"Working with: {logger.current_state['selected_file']}")
    
    # Add a separator
    st.markdown("---")

    # Initialize the analyzer and agent
    analyzer = JobDescriptionAnalyzer()
    agent = JobDescriptionAgent(model_id="anthropic.claude-3-haiku-20240307-v1:0")
    
    if 'final_version' not in st.session_state:
        st.session_state.final_version = None
    
    # Render the appropriate page based on current_page state
    if st.session_state.current_page == "generate":
        render_generation_page(logger, analyzer, agent)
    else:  # "refine"
        render_refinement_page(logger, analyzer, agent)
    
    # Footer with export options
    st.markdown("---")
    
    # Add simplified export options
    with st.expander("Export Options", expanded=False):
        st.subheader("Export Data")
        export_col1, export_col2 = st.columns(2)
        
        with export_col1:
            if st.button("Export All Feedback", key="export_all_feedback"):
                try:
                    # Get all sessions
                    all_sessions = JDOptimLogger.list_sessions()
                    
                    # Create data for CSV
                    csv_data = []
                    
                    for session_info in all_sessions:
                        session_id = session_info["session_id"]
                        log_file = os.path.join("logs", f"jdoptim_session_{session_id}.json")
                        
                        if os.path.exists(log_file):
                            with open(log_file, 'r') as f:
                                state = json.load(f)
                                
                                role = state.get("username", "Unknown")
                                start_time = state.get("session_start_time", "Unknown")
                                file_name = state.get("selected_file", "Unknown")
                                
                                # Add feedback items
                                for i, feedback in enumerate(state.get("feedback_history", [])):
                                    # Find feedback timestamp
                                    timestamp = start_time
                                    for action in state.get("actions", []):
                                        if action.get("action") == "feedback" and action.get("index", -1) == i:
                                            timestamp = action.get("timestamp", start_time)
                                            break
                                    
                                    # Parse feedback content
                                    if isinstance(feedback, dict):
                                        feedback_text = feedback.get("feedback", "")
                                        feedback_type = feedback.get("type", "General Feedback")
                                    else:
                                        feedback_text = feedback
                                        feedback_type = "General Feedback"
                                    
                                    # Add row to CSV data
                                    csv_data.append({
                                        "Session ID": session_id,
                                        "Role": role,
                                        "File": file_name,
                                        "Timestamp": timestamp,
                                        "Feedback Type": feedback_type,
                                        "Feedback": feedback_text
                                    })
                    
                    if csv_data:
                        # Convert to DataFrame and generate CSV
                        df = pd.DataFrame(csv_data)
                        csv = df.to_csv(index=False)
                        
                        # Offer download
                        st.download_button(
                            label="Download All Feedback CSV",
                            data=csv,
                            file_name=f"all_feedback_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                            mime="text/csv",
                            key="download_all_feedback_csv"
                        )
                    else:
                        st.warning("No feedback data found.")
                        
                except Exception as e:
                    st.error(f"Error exporting feedback: {str(e)}")
        
        with export_col2:
            if st.button("Export Current Session", key="export_session"):
                try:
                    # Get the current session data as JSON
                    session_data = json.dumps(logger.current_state, indent=2)
                    
                    # Offer download
                    st.download_button(
                        label="Download Session JSON",
                        data=session_data,
                        file_name=f"session_{logger.session_id}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        key="download_session_json"
                    )
                except Exception as e:
                    st.error(f"Error exporting session: {str(e)}")
    
    # Footer with company info
    footer_col1, footer_col2 = st.columns([4, 1])
    
    with footer_col1:
        st.caption("Job Description Enhancer | Made by Apexon")
    
    with footer_col2:
        st.caption("v2.0 - 2025")

if __name__ == "__main__":
    main()