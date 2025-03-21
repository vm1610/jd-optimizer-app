import os
from docx import Document

def read_job_description(file_path):
    """Read job description from either .txt or .docx file"""
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")

def save_enhanced_jd(content, filename, format_type):
    """Save job description content to a file"""
    if format_type == 'docx':
        doc = Document()
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(filename)
        return True
    elif format_type == 'txt':
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(content)
        return True
    return False

def get_jd_files():
    """Get a list of job description files from the JDs directory"""
    jd_directory = os.path.join(os.getcwd(), "JDs")
    if not os.path.exists(jd_directory):
        return []
    
    return [f for f in os.listdir(jd_directory) if f.endswith(('.txt', '.docx'))]

def get_feedback_files():
    """Get a list of feedback files from the Feedbacks directory"""
    feedback_directory = os.path.join(os.getcwd(), "Feedbacks")
    if not os.path.exists(feedback_directory):
        return []
    
    return [f for f in os.listdir(feedback_directory) if f.endswith(('.txt', '.docx'))]

def read_feedback_file(file_path):
    """Read feedback from a file"""
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")

def process_uploaded_docx(uploaded_file):
    """Process an uploaded docx file and return its content"""
    temp_path = f"temp_{uploaded_file.name}"
    
    try:
        # Save to temporary file
        with open(temp_path, 'wb') as f:
            f.write(uploaded_file.getvalue())
        
        # Read content using Document class
        doc = Document(temp_path)
        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        return content
    finally:
        # Clean up temp file
        if os.path.exists(temp_path):
            os.remove(temp_path)