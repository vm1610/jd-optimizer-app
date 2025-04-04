import streamlit as st
import os
import datetime
import tempfile
import json
from docx import Document
from ui.common import (
    display_section_header, display_subsection_header,
    display_warning_message, display_info_message, display_success_message,
    render_jd_selector, display_jd_comparison
)
from utils.file_utils import save_enhanced_jd

def render_client_feedback_page(services):
    """
    Render the Client Feedback page with seamless integration
    
    Args:
        services (dict): Dictionary of shared services
    """
    # Unpack services
    logger = services.get('logger')
    analyzer = services.get('analyzer')
    agent = services.get('agent')
    state_manager = services.get('state_manager')
    
    display_section_header("💬 Client Feedback Enhancement")
    
    # First, check if we have an active JD in our repository
    jd_repository = state_manager.get('jd_repository', {})
    jd_content, jd_source_name, jd_unique_id = state_manager.get_jd_content()
    
    if not jd_content:
        # No active JD, need to select one first
        display_subsection_header("1. Select Job Description")
        display_info_message("First, select a job description to enhance with client feedback.")
        
        # Use the unified JD selector
        has_jd = render_jd_selector(state_manager, services, "client_feedback")
        
        if not has_jd:
            st.warning("Please select a job description to continue.")
            return
        
        # Refresh JD content after selection
        jd_content, jd_source_name, jd_unique_id = state_manager.get_jd_content()
    else:
        # Show info about active JD
        st.success(f"Using: {jd_source_name}")
        
        # Option to select a different JD
        change_jd = st.checkbox("Use a different job description", value=False)
        if change_jd:
            display_subsection_header("1. Select Job Description")
            has_jd = render_jd_selector(state_manager, services, "client_feedback_alt")
            
            if has_jd:
                # Refresh JD content after selection
                jd_content, jd_source_name, jd_unique_id = state_manager.get_jd_content()
    
    # Create two columns for the main content
    feedback_col, preview_col = st.columns([1, 1])
    
    # Dictionary to store client feedback data
    client_feedback_data = state_manager.get('client_feedback_data', {
        'text': '',
        'type': 'Client Feedback',
        'upload_mode': 'Direct Entry'
    })
    
    # First column for feedback collection
    with feedback_col:
        display_subsection_header("2. Provide Client Feedback")
        
        # Define feedback types
        feedback_types = [
            "Client Feedback", 
            "Rejected Candidate Feedback", 
            "Hiring Manager Feedback", 
            "Selected Candidate Feedback", 
            "Interview Feedback"
        ]
        
        # Create tabs for feedback methods
        feedback_tabs = st.tabs(["Upload Feedback File", "Select from Directory", "Enter Manually"])
        
        # 1. Upload feedback file tab
        with feedback_tabs[0]:
            feedback_file = st.file_uploader(
                "📝 Drop or upload Feedback File",
                type=["txt", "docx", "csv"],
                key="client_feedback_upload",
                help="Upload the feedback from your client (TXT, DOCX, or CSV)"
            )
            
            # Feedback type selection
            selected_feedback_type = st.selectbox(
                "Feedback Type:",
                options=feedback_types,
                index=feedback_types.index(client_feedback_data.get('type', 'Client Feedback')) 
                    if client_feedback_data.get('type') in feedback_types else 0,
                key="client_feedback_type"
            )
            
            # Process uploaded feedback file
            if feedback_file:
                try:
                    if feedback_file.name.endswith(".txt"):
                        client_feedback = feedback_file.getvalue().decode("utf-8")
                    elif feedback_file.name.endswith(".csv"):
                        # Special handling for CSV
                        csv_content = feedback_file.getvalue().decode("utf-8")
                        client_feedback, message = process_csv_content(csv_content, "feedback")
                        st.info(message)
                    elif feedback_file.name.endswith(".docx"):
                        client_feedback = process_uploaded_docx(feedback_file)
                    else:
                        st.error("Unsupported file format.")
                        return
                    
                    # Update client feedback data
                    client_feedback_data['text'] = client_feedback
                    client_feedback_data['type'] = selected_feedback_type
                    client_feedback_data['upload_mode'] = 'File Upload'
                    state_manager.set('client_feedback_data', client_feedback_data)
                    
                    # Preview
                    with st.expander("Preview Feedback", expanded=True):
                        st.text_area("Feedback Content", client_feedback, height=200, disabled=True)
                        
                    display_success_message(f"Feedback loaded from {feedback_file.name}")
                except Exception as e:
                    st.error(f"Error reading feedback file: {str(e)}")
        
        # 2. Select from directory tab
        with feedback_tabs[1]:
            feedback_directory = os.path.join(os.getcwd(), "Data/Feedbacks")
            
            if not os.path.exists(feedback_directory):
                display_warning_message("The 'Feedbacks' directory does not exist. Create it or upload a file directly.")
            else:
                try:
                    feedback_files = [f for f in os.listdir(feedback_directory) 
                                    if f.endswith(('.txt', '.docx', '.csv'))]
                    
                    if not feedback_files:
                        display_warning_message("No feedback files found in the Feedbacks directory.")
                    else:
                        # Allow user to select a feedback file
                        selected_feedback_file = st.selectbox(
                            "Select Feedback File",
                            feedback_files,
                            help="Choose a feedback file to process",
                            key="folder_feedback_file"
                        )
                        
                        # Select feedback type
                        file_feedback_type = st.selectbox(
                            "Feedback Type:",
                            options=feedback_types,
                            index=feedback_types.index(client_feedback_data.get('type', 'Client Feedback'))
                                if client_feedback_data.get('type') in feedback_types else 0,
                            key="file_feedback_type"
                        )
                        
                        # Add a button to load the selected file
                        if st.button("Load Selected File", key="load_feedback_file"):
                            if selected_feedback_file:
                                feedback_path = os.path.join(feedback_directory, selected_feedback_file)
                                
                                if not os.path.exists(feedback_path):
                                    st.error(f"File not found: {feedback_path}")
                                else:
                                    try:
                                        if selected_feedback_file.endswith('.csv'):
                                            with open(feedback_path, 'r', encoding='utf-8') as f:
                                                csv_content = f.read()
                                            feedback_content, message = process_csv_content(csv_content, "feedback")
                                            st.info(message)
                                        else:
                                            feedback_content = read_feedback_file(feedback_path)
                                        
                                        # Update client feedback data
                                        client_feedback_data['text'] = feedback_content
                                        client_feedback_data['type'] = file_feedback_type
                                        client_feedback_data['upload_mode'] = 'Directory Selection'
                                        state_manager.set('client_feedback_data', client_feedback_data)
                                        
                                        # Display preview
                                        st.text_area(
                                            "Feedback Content",
                                            feedback_content,
                                            height=200,
                                            disabled=True,
                                            key="folder_feedback_content"
                                        )
                                        
                                        display_success_message(f"Successfully loaded feedback from {selected_feedback_file}")
                                    except Exception as e:
                                        st.error(f"Error reading feedback file: {str(e)}")
                except Exception as e:
                    st.error(f"Error accessing Feedbacks directory: {str(e)}")
        
        # 3. Manual input tab
        with feedback_tabs[2]:
            manual_feedback = st.text_area(
                "Enter client feedback:",
                value=client_feedback_data.get('text', '') if client_feedback_data.get('upload_mode') == 'Direct Entry' else '',
                height=200,
                placeholder="Enter the feedback from your client here...",
                key="manual_client_feedback"
            )
            
            manual_feedback_type = st.selectbox(
                "Feedback Type:",
                options=feedback_types,
                index=feedback_types.index(client_feedback_data.get('type', 'Client Feedback'))
                    if client_feedback_data.get('type') in feedback_types else 0,
                key="manual_feedback_type"
            )
            
            if st.button("Use This Feedback", key="use_manual_feedback"):
                if manual_feedback.strip():
                    # Update client feedback data
                    client_feedback_data['text'] = manual_feedback
                    client_feedback_data['type'] = manual_feedback_type
                    client_feedback_data['upload_mode'] = 'Direct Entry'
                    state_manager.set('client_feedback_data', client_feedback_data)
                    
                    display_success_message("Manual feedback saved!")
                else:
                    display_warning_message("Please enter some feedback first.")
    
    # Second column for preview
    with preview_col:
        display_subsection_header("Job Description Preview")
        with st.expander("View Current Job Description", expanded=True):
            st.text_area(
                "Current Content", 
                jd_content, 
                height=350, 
                disabled=True,
                key="client_jd_preview"
            )
            
        # Show feedback preview if available
        if client_feedback_data.get('text'):
            display_subsection_header("Feedback Preview")
            with st.expander("View Feedback", expanded=True):
                st.text_area(
                    f"{client_feedback_data.get('type')} Content", 
                    client_feedback_data.get('text'), 
                    height=200, 
                    disabled=True,
                    key="client_feedback_preview"
                )
    
    # --- Generate Enhanced JD Button ---
    st.markdown("---")
    display_subsection_header("3. Generate Enhanced Job Description")
    
    generate_col1, generate_col2 = st.columns([3, 1])
    
    with generate_col1:
        generate_btn = st.button(
            "🚀 Generate Enhanced Job Description", 
            type="primary", 
            key="generate_client_enhanced_jd",
            help="Generate an enhanced version of the job description based on client feedback"
        )
    
    with generate_col2:
        st.caption("AI will enhance the job description based on the provided client feedback.")
    
    # Check if we already have a client-enhanced version
    client_enhanced_jd = state_manager.get('client_enhanced_jd')
    
    # Handle generation process
    if generate_btn:
        if not jd_content:
            st.warning("Please provide a job description before generating.")
            return
            
        if not client_feedback_data.get('text'):
            st.warning("Please provide client feedback before generating.")
            return
        
        client_feedback = client_feedback_data.get('text')
        feedback_type = client_feedback_data.get('type')
        
        with st.spinner("Enhancing job description with client feedback..."):
            try:
                # Create feedback object with type
                feedback_obj = {
                    "feedback": client_feedback,
                    "type": feedback_type,
                    "role": state_manager.get('role')
                }
                
                # Add to feedback repository
                feedback_repository = state_manager.get('feedback_repository', {})
                history = feedback_repository.get('history', [])
                history.append(feedback_obj)
                state_manager.update_feedback_repository('history', history, source_tab="client_feedback")
                
                # Log feedback if logger is available
                if logger:
                    logger.current_state["feedback_history"].append(feedback_obj)
                    logger._save_state()
                    
                    # Log the action
                    logger.current_state["actions"].append({
                        "action": "client_feedback",
                        "timestamp": datetime.datetime.now().isoformat(),
                        "feedback_type": feedback_type
                    })
                    logger._save_state()
                
                # Create the prompt for the agent
                prompt = (
                    "You are an expert in job description refinement.\n\n"
                    "Please revise the provided job description **only based on the feedback** given by the client.\n\n"
                    "Do not introduce any information or changes not explicitly stated in the feedback.\n"
                    "Only make edits that directly reflect specific feedback content.\n\n"
                    "**Guidelines:**\n"
                    "- Do not make assumptions.\n"
                    "- Do not change formatting or structure unless required by feedback.\n"
                    "- Refer to the position as 'this role'.\n"
                    "- If the feedback is vague or irrelevant, leave the job description unchanged.\n\n"
                    f"### Original Job Description:\n{jd_content}\n\n"
                    f"### Client Feedback:\n{client_feedback}\n\n"
                    "### Please return only the revised job description below (leave unchanged if no edits are needed):\n"
                )
                
                # Call the agent to generate the enhanced JD
                enhanced_jd = agent.generate_final_description(jd_content, [feedback_obj])
                
                # Store the enhanced JD in state
                state_manager.set('client_enhanced_jd', enhanced_jd)
                
                # Log the enhanced version if logger is available
                if logger:
                    logger.log_enhanced_version(enhanced_jd, is_final=True)
                
                # Display success message
                display_success_message("Job description enhanced successfully based on client feedback!")
                
                # Force page refresh to show results
                st.rerun()
            except Exception as e:
                st.error(f"Error enhancing job description: {str(e)}")
                st.error("Please try again or contact support if the problem persists.")
    
    # --- Display Results ---
    if client_enhanced_jd:
        # Display results in an organized layout
        st.markdown("---")
        display_section_header("Results")
        
        # Use the comparison component
        display_jd_comparison(jd_content, client_enhanced_jd, services, "client_feedback")
        
        # Download options
        display_section_header("Download Options")
        
        download_col1, download_col2 = st.columns(2)
        
        with download_col1:
            st.download_button(
                label="Download as TXT",
                data=client_enhanced_jd,
                file_name=f"client_enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                key="client_download_txt"
            )
            if logger:
                logger.log_download("txt", f"client_enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
        
        with download_col2:
            if st.button("Download as DOCX", key="client_download_docx"):
                docx_filename = f"client_enhanced_jd_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                save_enhanced_jd(client_enhanced_jd, docx_filename, 'docx')
                display_success_message(f"Saved as {docx_filename}")
                if logger:
                    logger.log_download("docx", docx_filename)
                    
def read_feedback_file(file_path):
    """Read feedback from a file"""
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_path.endswith('.csv'):
        # Read CSV as raw text first
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    else:
        raise ValueError("Unsupported file format")

def process_uploaded_docx(uploaded_file):
    """Process an uploaded docx file and return its content"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(uploaded_file.getvalue())
        temp_path = tmp.name
    
    try:
        doc = Document(temp_path)
        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return content
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

def process_csv_content(csv_content, column_type="feedback"):
    """Extract content from CSV files based on column type"""
    try:
        import pandas as pd
        df = pd.read_csv(pd.StringIO(csv_content))
        
        # Define potential column names based on content type
        if column_type == "feedback":
            potential_columns = ['feedback', 'comments', 'notes', 'review', 'suggestions', 'input']
        else:  # job description
            potential_columns = ['job_description', 'description', 'jd', 'content', 'text']
        
        # Find the first matching column or use the first text column
        target_column = None
        for col in potential_columns:
            if col in df.columns:
                target_column = col
                break
        
        if target_column is None:
            # If no matching column found, use the first column that seems to have text content
            for col in df.columns:
                if df[col].dtype == 'object' and df[col].str.len().mean() > (50 if column_type == "jd" else 20):
                    target_column = col
                    break
        
        if target_column:
            if column_type == "feedback":
                # Combine all feedback entries
                combined_content = "\n\n".join(df[target_column].dropna().tolist())
                return combined_content, f"Extracted {len(df[target_column].dropna())} entries from column: {target_column}"
            else:
                # For JD, just use the first non-empty value
                return df[target_column].dropna().iloc[0], f"Extracted job description from column: {target_column}"
        
        # Fallback - concatenate all text columns
        text_cols = [col for col in df.columns if df[col].dtype == 'object']
        combined_content = "\n\n".join([f"{col}:\n{df[col].iloc[0]}" for col in text_cols[:5]])
        return combined_content, f"Could not identify a specific {column_type} column. Using combined text from CSV."
    except Exception as e:
        return csv_content, f"Error processing CSV structure: {str(e)}. Using raw CSV content."